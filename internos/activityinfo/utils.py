#!/usr/bin/env python
# coding=utf-8
# -*- coding: utf-8 -*-

import os
import csv
import datetime
import subprocess
import logging
import calendar
from datetime import date

from django.db.models import Sum, Q
from internos.activityinfo.client import Client
from internos.activityinfo.exports import get_database_data, read_file, get_xlsx
from .utils_common import *


logger = logging.getLogger(__name__)


from collections import defaultdict
from django.apps import apps


class BulkCreateManager(object):
    """
    This helper class keeps track of ORM objects to be created for multiple
    model classes, and automatically creates those objects with `bulk_create`
    when the number of objects accumulated for a given model class exceeds
    `chunk_size`.
    Upon completion of the loop that's `add()`ing objects, the developer must
    call `done()` to ensure the final set of objects is created for all models.
    """

    def __init__(self, chunk_size=100):
        self._create_queues = defaultdict(list)
        self.chunk_size = chunk_size

    def _commit(self, model_class):
        model_key = model_class._meta.label
        model_class.objects.bulk_create(self._create_queues[model_key])
        self._create_queues[model_key] = []

    def add(self, obj):
        """
        Add an object to the queue to be created, and call bulk_create if we
        have enough objs.
        """
        model_class = type(obj)
        model_key = model_class._meta.label
        self._create_queues[model_key].append(obj)
        if len(self._create_queues[model_key]) >= self.chunk_size:
            self._commit(model_class)

    def done(self):
        """
        Always call this upon completion to make sure the final partial chunk
        is saved.
        """
        for model_name, objs in self._create_queues.items():
            if len(objs) > 0:
                self._commit(apps.get_model(model_name))


def get_extraction_month(ai_db):

    current_year = date.today().year
    current_month = date.today().month
    reporting_year = ai_db.reporting_year.year

    if current_year - 1 >= int(reporting_year):
        return 13
    else:
        return current_month


def get_current_extraction_month(ai_db):
    current_year = date.today().year
    current_month = date.today().month
    reporting_year = ai_db.reporting_year.year

    if current_year - 1 == int(reporting_year) and current_month == 1:
        return 12
    else:
        return current_month


def get_live_extraction_month(ai_db):
    current_year = date.today().year
    current_month = date.today().month
    reporting_year = ai_db.reporting_year.year

    if current_year - 1 == int(reporting_year) and current_month == 1:
        return 12
    else:
        return current_month


def r_script_command_line_old(script_name, ai_db):
    command = 'Rscript'
   # path = os.path.dirname(os.path.abspath(__file__))
    path = os.path.dirname(__file__)
    path2script = os.path.join(path, 'RScripts')
    path2script = os.path.join(path2script, script_name)
    main_db_id = ai_db.reporting_year.database_id
    cmd = [command, path2script, ai_db.username, ai_db.password, str(ai_db.db_id), str(ai_db.ai_id), str(main_db_id)]
    try:
         subprocess.check_output(cmd, universal_newlines=True)
        # subprocess.check_output(['type', cmd], shell=True)
    except subprocess.CalledProcessError as e:
        print("command '{}' return with error (code {}): {}".format(e.cmd, e.returncode, e.output))
        return 0

    return 1


def r_script_command_line(ai_db):

    client = Client()
    filters = "LEFT(Month,4) == '{}'".format(ai_db.reporting_year.year)
    if ai_db.parent_id:
        get_database_data(client, database_id=ai_db.parent_id, resource_id=ai_db.db_id, database=ai_db,
                          record_filter=filters)
        get_xlsx(client, database_id=ai_db.parent_id, resource_id=ai_db.db_id, database=ai_db, record_filter=filters)
    else:
        get_database_data(client, database_id=ai_db.db_id, database=ai_db, record_filter=filters)


def read_data_from_file(ai_db, forced=False, report_type=None):
    from internos.activityinfo.models import ActivityReport, LiveActivityReport
    result = 0

    if report_type == 'live':
        model = LiveActivityReport.objects.none()
        LiveActivityReport.objects.filter(database_id=ai_db.ai_id).delete()
        result = add_rows(ai_db=ai_db, model=model)

    if forced:
        model = ActivityReport.objects.none()
        ActivityReport.objects.filter(database_id=ai_db.ai_id).delete()
        if ai_db.have_offices:
            result = add_rows_temp(ai_db=ai_db, model=model)
        else:
            result = add_rows(ai_db=ai_db, model=model)

    # link_ai_partners(report_type)
    return result


def import_data_via_r_script(ai_db, report_type=None):
    r_script_command_line(ai_db)
    total = read_data_from_file(ai_db, True, report_type)
    return total


def get_awp_code(name):
    try:
        if '_' in name:
            awp_code = name[:name.find('_')]
        elif ' - ' in name:
            awp_code = name[:name.find(' - ')]
            if ': ' in awp_code:
                awp_code = awp_code[:awp_code.find(': ')]
        elif ': ' in name:
            awp_code = name[:name.find(': ')]
        else:
            awp_code = name[:name.find('#')]
            if ': ' in awp_code:
                awp_code = awp_code[:awp_code.find(': ')]
    except TypeError as ex:
        awp_code = 'None'
    return awp_code


def get_label(data):
    try:
        if '------' in data['name']:
            return data['description']
    except TypeError as ex:
        pass
    return data['name']


def set_tags(indicator, tags):
    for tag in tags:
        if tag.name in indicator.name or tag.name.title() in indicator.name:
            setattr(indicator, tag.tag_field, tag)
    indicator.save()


def clean_string(value, string):
    return value.replace(string, '')


def add_rows_temp(ai_db=None, model=None):

    month = get_current_extraction_month(ai_db)
    path = os.path.dirname(os.path.abspath(__file__))
    path2file = path+'/AIReports/'+str(ai_db.ai_id)+'_ai_data.txt'
    values = read_file(path2file)
    ctr = 0

    for row in values:
        indicator_value = 0
        if 'Value' in row:
            indicator_value = row['Value']

        try:
            indicator_value = float(indicator_value)
        except Exception:
            indicator_value = 0

        funded_by = 'UNICEF'
        partner_label = 'UNICEF'

        start_date = None
        if 'month' in row and row['month'] and not row['month'] == 'NA':
            start_date = '{}-01'.format(row['month'])

        if ai_db.reporting_year.year not in start_date:
            continue

        gov_code = row['reporting_office']
        gov_name = row['reporting_office']

        model.create(
            month=month,
            database=row['Folder'],
            database_id=ai_db.ai_id,
            report_id=row['FormId'],
            indicator_id=row['Quantity Field ID'],
            indicator_name=row['Quantity Field'],
            indicator_awp_code='',
            month_name=row['month'] if 'month' in row else '',
            partner_label=partner_label,
            location_adminlevel_caza_code=row['caza.code'] if 'caza.code' in row else '',
            location_adminlevel_caza=row['caza.name'] if 'caza.name' in row else '',
            form=row['Form'] if 'Form' in row else '',
            location_adminlevel_cadastral_area_code=row['cadastral_area.code'] if 'cadastral_area.code' in row else'',
            location_adminlevel_cadastral_area=row['cadastral_area.name'] if 'cadastral_area.name' in row else '',
            governorate=row['governorate'] if 'governorate' in row else '',
            location_adminlevel_governorate_code=gov_code,
            location_adminlevel_governorate=gov_name,
            partner_description='UNICEF',
            project_start_date=row['projects.start_date'] if 'projects.start_date' in row and not row['projects.start_date'] == 'NA' else None,
            project_end_date=row['projects.end_date'] if 'projects.end_date' in row and not row['projects.start_date'] == 'NA' else None,
            project_label=row['projects.project_code'] if 'projects.project_code' in row else '',
            project_description=row['projects.project_name'] if 'projects.project_name' in row else '',
            funded_by=funded_by,
            indicator_value=indicator_value,
            indicator_units=row['units'] if 'units' in row else '',
            reporting_section=row['reporting_section'] if 'reporting_section' in row else '',
            site_type=row['site_type'] if 'site_type' in row else '',
            location_longitude=row[
                'ai_allsites.geographic_location.longitude'] if 'ai_allsites.geographic_location.longitude' in row else '',
            location_latitude=row[
                'ai_allsites.geographic_location.latitude'] if 'ai_allsites.geographic_location.latitude' in row else '',
            location_alternate_name=row[
                'ai_allsites.alternate_name'] if 'ai_allsites.alternate_name' in row else '',

            location_name=row['ai_allsites.name'] if 'ai_allsites.name' in row else '',
            partner_id=row['partner_id'] if 'partner_id' in row else partner_label,

            start_date=start_date,
        )
        ctr += 1

    return ctr


def add_rows(ai_db=None, model=None):
    from internos.activityinfo.models import ActivityReport, LiveActivityReport
    month = get_current_extraction_month(ai_db)
    path = os.path.dirname(os.path.abspath(__file__))
    path2file = path+'/AIReports/'+str(ai_db.ai_id)+'_ai_data.txt'
    values = read_file(path2file)
    ctr = 0

    print('start add rows')

    bulk_mgr = BulkCreateManager(chunk_size=20)

    for row in values:
        indicator_value = 0
        if 'Value' in row:
            indicator_value = row['Value']

        try:
            indicator_value = float(indicator_value)
        except Exception:
            indicator_value = 0

        funded_by = row['funded_by.funded_by'] if 'funded_by.funded_by' in row else ''
        partner_label = row['partner.name'] if 'partner.name' in row else ''
        partner_label = partner_label.replace('-', '_')
        partner_label = partner_label.encode("utf-8").replace('é', 'e')

        if partner_label == 'UNICEF':
            funded_by = 'UNICEF'

        start_date = None
        if 'month' in row and row['month'] and not row['month'] == 'NA':
            date = row['month']
            try:
                date = datetime.datetime.strptime(date, '%Y-%m-%d')
                start_date = date
            except Exception:
                 start_date = '{}-01'.format(date)

        if 'month.' in row and row['month.'] and not row['month.'] == 'NA':
            start_date = '{}-01'.format(row['month.'])
        if 'date' in row and row['date'] and not row['date'] == 'NA':
            start_date = row['date']

        gov_code = 0
        gov_name = ""
        if 'governorate.code' in row:
            if row['governorate.code'] == 'NA':
                gov_code = 10
            else:
                gov_code = row['governorate.code']

        if 'governorate.name' in row:
            if row['governorate.name'] == 'NA':
                gov_name = "National"
            else:
                gov_name = row['governorate.name']

        support_covid1 = False
        support_covid2 = False
        support_covid3 = False

        if 'X4.2.3_covid_adaptation' in row:
            if row['X4.2.3_covid_adaptation'] == 'Yes':
                support_covid1 = True

        if 'covid_adaptation' in row:
            if row['covid_adaptation'] == 'Yes':
                support_covid2 = True

        if 'covid_adapted_sensitization' in row:
            if row['covid_adapted_sensitization'] == 'Yes':
                support_covid3 = True

        support_covid = support_covid1 or support_covid2 or support_covid3
        print(start_date)

        bulk_mgr.add(ActivityReport(
            month=month,
            database=row['Folder'],
            database_id=ai_db.ai_id,
            report_id=row['FormId'],
            indicator_id=row['Quantity Field ID'],
            indicator_name=row['Quantity Field'],
            indicator_awp_code=get_awp_code(row['Quantity Field']),
            month_name=row['month'] if 'month' in row else '',
            partner_label=partner_label,
            location_adminlevel_caza_code=row['caza.code'] if 'caza.code' in row else '',
            location_adminlevel_caza=row['caza.name'] if 'caza.name' in row else '',
            form=row['Form'] if 'Form' in row else '',
            location_adminlevel_cadastral_area_code=row['cadastral_area.code'] if 'cadastral_area.code' in row else'',
            location_adminlevel_cadastral_area=row['cadastral_area.name'] if 'cadastral_area.name' in row else '',
            governorate=row['governorate'] if 'governorate' in row else '',
            location_adminlevel_governorate_code=gov_code,
            location_adminlevel_governorate=gov_name,
            partner_description=row['partner.partner_full_name'] if 'partner.partner_full_name' in row else '',
            project_start_date=row['projects.start_date'] if 'projects.start_date' in row and not row['projects.start_date'] == 'NA' else None,
            project_end_date=row['projects.end_date'] if 'projects.end_date' in row and not row['projects.start_date'] == 'NA' else None,
            project_label=row['projects.project_code'] if 'projects.project_code' in row else '',
            project_description=row['projects.project_name'] if 'projects.project_name' in row else '',
            funded_by=funded_by,
            indicator_value=indicator_value,
            indicator_units=row['units'] if 'units' in row else '',
            reporting_section=row['reporting_section'] if 'reporting_section' in row else '',
            site_type=row['site_type'] if 'site_type' in row else '',
            location_longitude=row[
                'ai_allsites.geographic_location.longitude'] if 'ai_allsites.geographic_location.longitude' in row else '',
            location_latitude=row[
                'ai_allsites.geographic_location.latitude'] if 'ai_allsites.geographic_location.latitude' in row else '',
            location_alternate_name=row[
                'ai_allsites.alternate_name'] if 'ai_allsites.alternate_name' in row else '',
            location_name=row['ai_allsites.name'] if 'ai_allsites.name' in row else '',
            partner_id=row['partner_id'] if 'partner_id' in row else partner_label,
            support_covid=support_covid,
            start_date=start_date,
        ))
        ctr += 1

    print('end add rows')
    bulk_mgr.done()
    return ctr


def update_rows(ai_id):
    from internos.activityinfo.models import ActivityReport
    path = os.path.dirname(os.path.abspath(__file__))
    path2file = path+'/AIReports/'+str(ai_id)+'_ai_data.csv'

    with open(path2file) as csvfile:
        reader = csv.DictReader(csvfile)
        ctr = 0
        for row in reader:
            ctr += 1
            ActivityReport.objects.update(
                database_id=ai_id,
                indicator_id=clean_string(row['indicator.id'], 'i'),
                indicator_units=row['indicator.units'] if 'indicator.units' in row else '',
                indicator_value=row['indicator.value'] if 'indicator.value' in row else '',
            )
        return ctr


def generate_indicator_awp_code(ai_id):
    from internos.activityinfo.models import Indicator

    data = Indicator.objects.filter(activity__database__ai_id=ai_id)
    ctr = data.count()
    for item in data:
        item.awp_code = get_awp_code(item.name)
        item.save()

    return ctr


def generate_indicator_tag(ai_id):
    from internos.activityinfo.models import Indicator, IndicatorTag

    data = Indicator.objects.filter(activity__database__ai_id=ai_id)
    for item in data :
        item.tag_age = None
        item.tag_nationality= None
        item.tag_disability = None
        item.tag_programme = None
        item.tag_gender =  None
        item.save()
    data = Indicator.objects.filter(activity__database__ai_id=ai_id ,master_indicator=False,master_indicator_sub =False)
    tags = IndicatorTag.objects.all()

    ctr = data.count()
    for item in data:
        set_tags(item, tags)

    return ctr


def generate_indicator_awp_code2(ai_id):
    from internos.activityinfo.models import ActivityReport

    data = ActivityReport.objects.filter(database_id=ai_id)
    ctr = data.count()
    for item in data:
        item.indicator_awp_code = get_awp_code(item.indicator_name)
        item.save()

    return ctr


def calculate_sum_target(ai_id):
    from internos.activityinfo.models import Indicator

    top_indicators = Indicator.objects.filter(master_indicator=True, activity__database_id=ai_id)
    sub_indicators = Indicator.objects.filter(master_indicator_sub=True, activity__database_id=ai_id)

    for item in sub_indicators:
        if not item.summation_sub_indicators.count():
            continue
        target_sum = item.summation_sub_indicators.exclude(master_indicator=True).aggregate(Sum('target'))
        item.target = target_sum['target__sum'] if target_sum['target__sum'] else 0
        item.save()

    for item in top_indicators:
        target_sum = item.summation_sub_indicators.exclude(master_indicator_sub=False,
                                                           master_indicator=False).aggregate(Sum('target'))
        item.target = target_sum['target__sum'] if target_sum['target__sum'] else 0

        item.save()

    return top_indicators.count() + sub_indicators.count()


def generate_indicators_number(ai_db):
    from internos.activityinfo.models import Indicator

    for indicator in Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id):
        indicator.ai_indicator = indicator.get_ai_indicator
        indicator.save()


def link_indicators_data(ai_db, report_type=None):
    result = 0
    result = link_indicators_activity_report(ai_db, report_type)
    # link_indicators_project(ai_db)
    # link_etools_partnerships(ai_db)

    return result


#  todo for review
def link_indicators_project(ai_db):
    from internos.activityinfo.models import Indicator
    from internos.etools.models import PCA

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                          master_indicator=True).exclude(project_code__isnull=True).only(
        'project'
    )

    for item in indicators.iterator():
        try:
            project = PCA.objects.get(reference_number=item.project_code)
            item.project = project
            item.save()
        except Exception:
            continue


def link_indicators_activity_report(ai_db, report_type=None):
    from internos.activityinfo.models import Indicator, ActivityReport, LiveActivityReport

    ctr = 0
    if report_type == 'live':
        reports = LiveActivityReport.objects.filter(database_id=ai_db.ai_id)
    else:
        reports = ActivityReport.objects.filter(database_id=ai_db.ai_id)

    reports = reports.exclude(ai_indicator__isnull=False)

    # if ai_db.is_funded_by_unicef:
    #     reports = reports.filter(funded_by='UNICEF')

    indicators = Indicator.objects.filter(
        activity__database__ai_id=ai_db.ai_id).exclude(
        master_indicator=True).exclude(master_indicator_sub=True)

    for item in indicators:
        if not item.ai_indicator:
            continue
        ai_values = reports.filter(indicator_id=item.ai_indicator)

        if not ai_values.count():
            continue
        ctr += ai_values.count()
        ai_values.update(ai_indicator_id=item.id)

        try:
            item.project_code = ai_values.first().project_label
            item.project_name = ai_values.first().project_description

            item.save()

            main_master_indicator = item.main_master_indicator
            if main_master_indicator:
                main_master_indicator.project_code = item.project_code
                main_master_indicator.project_name = item.project_name
                main_master_indicator.save()
        except Exception:
            pass

    return ctr


def link_ai_locations(report_type=None):
    from internos.activityinfo.models import AdminLevelEntities, ActivityReport, LiveActivityReport

    ctr = 0
    if report_type == 'live':
        indicators = LiveActivityReport.objects.all()
    else:
        indicators = ActivityReport.objects.all()

    location_cadastral = indicators.values('location_adminlevel_cadastral_area_code').distinct()
    location_caza = indicators.values('location_adminlevel_caza_code').distinct()
    location_governorate = indicators.values('location_adminlevel_governorate_code').distinct()

    for item in location_cadastral:
        ai_values = indicators.filter(location_adminlevel_cadastral_area_code=item['location_adminlevel_cadastral_area_code'])

        if not ai_values.count():
            continue
        ctr += ai_values.count()
        try:
            ai_values.update(location_cadastral=AdminLevelEntities.objects.get(code=item['location_adminlevel_cadastral_area_code']))
        except Exception as ex:
            print(item['location_adminlevel_cadastral_area_code'])
            pass

    for item in location_caza:
        ai_values = indicators.filter(location_adminlevel_caza_code=item['location_adminlevel_caza_code'])

        if not ai_values.count():
            continue
        ctr += ai_values.count()
        try:
            ai_values.update(location_caza=AdminLevelEntities.objects.get(code=item['location_adminlevel_caza_code']))
        except Exception as ex:
            print(item['location_adminlevel_caza_code'])
            pass

    for item in location_governorate:
        ai_values = indicators.filter(location_adminlevel_governorate_code=item['location_adminlevel_governorate_code'])

        if not ai_values.count():
            continue
        ctr += ai_values.count()
        try:
            ai_values.update(location_governorate=AdminLevelEntities.objects.get(code=item['location_adminlevel_governorate_code']))
        except Exception as ex:
            print(item['location_adminlevel_governorate_code'])
            pass

    return ctr


def link_ai_partners(report_type=None):
    from internos.activityinfo.models import Partner, ActivityReport, LiveActivityReport

    ctr = 0
    if report_type == 'live':
        reports = LiveActivityReport.objects.all()
    else:
        reports = ActivityReport.objects.all()

    partners = Partner.objects.all()

    for item in partners:
        ai_values = reports.filter(partner_id=item.ai_number)
        if not ai_values.count():
            continue
        ctr += ai_values.count()
        ai_values.update(partner_ai_id=item.id)
        item.number = item.ai_number
        item.save()

    return ctr


def link_etools_partners():
    from internos.activityinfo.models import Partner
    from internos.etools.models import PartnerOrganization

    partners = PartnerOrganization.objects.all()

    for partner in partners:
        ai_partners = Partner.objects.filter(Q(name=partner.short_name) |
                                             Q(name=partner.name.upper()) |
                                             Q(name=partner.short_name.upper()) |
                                             Q(full_name=partner.short_name) |
                                             Q(name=partner.name) |
                                             Q(full_name=partner.name))
        if not ai_partners.count():
            continue

        ai_partners.update(partner_etools=partner)


#  todo for review
def link_etools_partnerships(ai_db=None):
    from internos.activityinfo.models import ActivityReport
    from internos.etools.models import PCA

    programmes = PCA.objects.all()

    ai_reports = ActivityReport.objects.filter(project_label__isnull=False)
    if ai_db:
        ai_reports = ai_reports.filter(database_id=ai_db.ai_id)

    for programme in programmes:
        reports = ai_reports.filter(project_label=programme.number)

        if not reports.count():
            continue

        reports.update(programme_document=programme)


def reset_indicators_values(ai_id, report_type=None):
    from internos.activityinfo.models import Indicator

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_id)
    for indicator in indicators:
        if report_type == 'live':
            indicator.values_live = {}
            indicator.values_gov_live = {}
            indicator.values_partners_live = {}
            indicator.values_partners_gov_live = {}
            indicator.cumulative_values_live = {}
            indicator.values_sections_live = {}
            indicator.values_sections_partners_live = {}
            indicator.values_sections_gov_live = {}
            indicator.values_sections_partners_gov_live = {}

        else:
            indicator.values = {}
            indicator.values_gov = {}
            indicator.values_partners = {}
            indicator.values_partners_gov = {}
            indicator.values_sections = {}
            indicator.values_sections_partners={}
            indicator.values_sections_gov={}
            indicator.values_sections_partners_gov={}
            indicator.cumulative_values = {}
        indicator.save()

    return indicators.count()


def calculate_indicators_values(ai_db, report_type=None):

    print('reset_indicators_values')
    reset_indicators_values(ai_db.ai_id, report_type)
    print('calculate_individual_indicators_values_2')
    if report_type == 'live':
        calculate_individual_indicators_values_2(ai_db)
    else:
        calculate_individual_indicators_values_1(ai_db)
    print('calculate_master_indicators_values_1')
    calculate_master_indicators_values_1(ai_db, report_type, True)
    print('calculate_master_indicators_values_1')
    calculate_master_indicators_values_1(ai_db, report_type)
    print('calculate_master_indicators_values_percentage')
    calculate_master_indicators_values_percentage(ai_db, report_type)
    print('calculate_master_indicators_values_denominator_multiplication')
    calculate_master_indicators_values_denominator_multiplication(ai_db, report_type)
    print('calculate_indicators_values_percentage_1')
    calculate_indicators_values_percentage_1(ai_db, report_type)  # DONE
    print('calculate_indicators_cumulative_results_1')
    calculate_indicators_cumulative_results_1(ai_db, report_type)
    print('calculate_indicators_highest_results_1')
    calculate_indicators_highest_results(ai_db ,report_type)
    print('calculate_indicators_status')
    calculate_indicators_status(ai_db)
    if report_type != 'live':
        print('calculate_indicators_all_tags')
        calculate_indicators_all_tags(ai_db)
    return 0


def calculate_indicators_cumulative_results_1(ai_db, report_type=None):
    from django.db import connection
    from internos.activityinfo.models import Indicator

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id).only(
        'id',
        'values',
        'values_gov',
        'values_partners',
        'values_partners_gov',
        'values_live',
        'values_gov_live',
        'values_partners_live',
        'values_partners_gov_live',
        'values_hpm',
        'cumulative_values',
        'cumulative_values_live',
        'values_sections',
        'values_sections_partners',
        'values_sections_gov',
        'values_sections_partners_gov',
        'values_sections_live',
        'values_sections_partners_live',
        'values_sections_gov_live',
        'values_sections_partners_gov_live'

    )

    rows_data = {}
    cursor = connection.cursor()
    cursor.execute(
        "SELECT distinct ai.id, ai.ai_indicator, aa.id, aa.name, ai.values, ai.values_live, "
        "ai.values_gov, ai.values_gov_live, ai.values_partners, ai.values_partners_live, "
        "ai.values_partners_gov, ai.values_partners_gov_live , ai.values_sections, ai.values_sections_live, "
        "ai.values_sections_partners, ai.values_sections_partners_live, ai.values_sections_gov, "
        "ai.values_sections_gov_live, ai.values_sections_partners_gov, "
        "ai.values_sections_partners_gov_live "
        "FROM public.activityinfo_indicator ai, public.activityinfo_activity aa "
        "WHERE ai.activity_id = aa.id AND aa.database_id = %s",
        [ai_db.id])

    rows = cursor.fetchall()
    for row in rows:
        rows_data[row[0]] = row

    for indicator in indicators.iterator():
        values_month = {}
        values_partners = {}
        values_gov = {}
        values_partners_gov = {}
        values_sections = {}
        values_sections_partners = {}
        values_sections_gov = {}
        values_sections_partners_gov = {}

        if indicator.id in rows_data:
            indicator_values = rows_data[indicator.id]

            if report_type == 'live':
                values = indicator_values[5]  # values_live
                values1 = indicator_values[7]  # values_gov_live
                values2 = indicator_values[9]  # values_partners_live
                values3 = indicator_values[11]  # values_partners_gov_live
                values4 = indicator_values[13]  # values_sections_live
                values5 = indicator_values[15]  # values_sections_partners_live
                values6 = indicator_values[17]  # values_sections_gov_live
                values7 = indicator_values[19]  # values_sections_partners_gov_live
            else:
                values = indicator_values[4]  # values
                values1 = indicator_values[6]  # values_gov
                values2 = indicator_values[8]  # values_partners
                values3 = indicator_values[10]  # values_partners_gov
                values4 = indicator_values[12]  # values_sections
                values5 = indicator_values[14]  # values_sections_partners
                values6 = indicator_values[16]  # values_sections_gov
                values7 = indicator_values[18]  # values_sections_partners_gov

            # for month in sorted(values):
            #     c_value = 0
            #     for c_month in range(1, int(month) + 1):
            #         if str(c_month) in values:
            #             c_value += float(values[str(c_month)])
            #         values_month[str(month)] = c_value

            c_value = 0
            for key, value in values.items():
                c_value += value
                values_month = c_value

            # for key in sorted(values1):
            #     c_value = 0
            #     for c_key in range(0, int(sorted(values1.keys()).index(key)) + 1):
            #         c_key = sorted(values1)[c_key]
            #         c_value += float(values1[c_key])
            #         values_gov[c_key] = c_value

            for key, value in values1.items():

                keys = key.split('-')
                gov = keys[1]
                if gov in values_gov:
                    values_gov[gov] = values_gov[gov] + value
                else:
                    values_gov[gov] = value

            # for key in sorted(values2):
            #     c_value = 0
            #     for c_key in range(0, int(sorted(values2.keys()).index(key)) + 1):
            #         c_key = sorted(values2)[c_key]
            #         c_value += float(values2[c_key])
            #         values_partners[c_key] = c_value

            for key, value in values2.items():
                keys = key.split('-')
                partner = keys[1]
                if partner in values_partners:
                    values_partners[partner] = values_partners[partner] + value
                else:
                    values_partners[partner] = value

            # for key in sorted(values3):
            #     c_value = 0
            #     for c_key in range(0, int(sorted(values3.keys()).index(key)) + 1):
            #         c_key = sorted(values3)[c_key]
            #         c_value += float(values3[c_key])
            #         values_partners_gov[c_key] = c_value

            for key, value in values3.items():
                keys = key.split('-')
                gov_partner = '{}-{}'.format(keys[1], keys[2])
                if gov_partner in values_partners_gov:
                    values_partners_gov[gov_partner] = values_partners_gov[gov_partner] + value
                else:
                    values_partners_gov[gov_partner] = value

            for key, value in values4.items():
                keys = key.split('-')
                section = keys[1]
                if section in values_sections:
                    values_sections[section] = values_sections[section] + value
                else:
                    values_sections[section] = value

            for key, value in values5.items():
                keys = key.split('-')
                section_partner = '{}-{}'.format(keys[1], keys[2])
                if section_partner in values_sections_partners:
                    values_sections_partners[section_partner] = values_sections_partners[section_partner] + value
                else:
                    values_sections_partners[section_partner] = value

            for key, value in values6.items():
                keys = key.split('-')
                section_gov = '{}-{}'.format(keys[1], keys[2])
                if section_gov in values_sections_gov:
                    values_sections_gov[section_gov] = values_sections_gov[section_gov] + value
                else:
                    values_sections_gov[section_gov] = value

            for key, value in values7.items():
                keys = key.split('-')
                section_partner_gov = '{}-{}-{}'.format(keys[1], keys[2],keys[3])
                if section_partner_gov in values_sections_partners_gov:
                    values_sections_partners_gov[section_partner_gov] = values_sections_partners_gov[section_partner_gov] + value
                else:
                    values_sections_partners_gov[section_partner_gov] = value

            if report_type == 'live':
                indicator.cumulative_values_live = {
                    'months': values_month,
                    'partners': values_partners,
                    'govs': values_gov,
                    'partners_govs': values_partners_gov,
                    'sections': values_sections,
                    'sections_partners':values_sections_partners,
                    'sections_gov':values_sections_gov,
                    'sections_partners_gov':values_sections_partners_gov
                }
            else:
                indicator.cumulative_values = {
                    'months': values_month,
                    'partners': values_partners,
                    'govs': values_gov,
                    'partners_govs': values_partners_gov,
                    'sections': values_sections,
                    'sections_partners': values_sections_partners,
                    'sections_gov': values_sections_gov,
                    'sections_partners_gov': values_sections_partners_gov
                }

            indicator.save()

    return indicators.count()


def calculate_indicators_cumulative_results(ai_db, report_type=None):
    from internos.activityinfo.models import Indicator, ActivityReport, LiveActivityReport

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id)

    if report_type == 'live':
        report = LiveActivityReport.objects.filter(database=ai_db)
    else:
        report = ActivityReport.objects.filter(database=ai_db)

    if ai_db.is_funded_by_unicef:
        report = report.filter(funded_by__contains='UNICEF')

    partners = report.values('partner_id').order_by('partner_id').distinct('partner_id')
    governorates = report.values('location_adminlevel_governorate_code').order_by('location_adminlevel_governorate_code').distinct('location_adminlevel_governorate_code')

    for indicator in indicators:
        value = 0
        cum_month = {}
        cum_partner = {}
        cum_gov = {}
        cum_partner_gov = {}

        if report_type == 'live':
            values = indicator.values_live
            values_partners = indicator.values_partners_live
            values_govs = indicator.values_gov_live
            values_partners_govs = indicator.values_partners_gov_live
        else:
            values = indicator.values
            values_partners = indicator.values_partners
            values_govs = indicator.values_gov
            values_partners_govs = indicator.values_partners_gov

        for month in values:
            c_value = 0
            for c_month in range(1, int(month) + 1):
                c_value = 0
                if c_month in values:
                    c_value += float(values[str(c_month)])
                cum_month[str(month)] = c_value

            for partner in partners:
                key = '{}-{}'.format(month, partner['partner_id'])
                if key in values_partners:
                    cum_partner[partner['partner_id']] = values_partners[key] + (cum_partner[partner['partner_id']] if partner['partner_id'] in cum_partner else 0)

            for gov in governorates:
                key = '{}-{}'.format(month, gov['location_adminlevel_governorate_code'])
                if key in values_govs:
                    cum_gov[gov['location_adminlevel_governorate_code']] = values_govs[key] + (cum_gov[gov['location_adminlevel_governorate_code']] if gov['location_adminlevel_governorate_code'] in cum_gov else 0)

                for partner in partners:
                    key = '{}-{}-{}'.format(month, partner['partner_id'], gov['location_adminlevel_governorate_code'])
                    key_c = '{}-{}'.format(partner['partner_id'], gov['location_adminlevel_governorate_code'])
                    if key in values_partners_govs:
                        cum_partner_gov[key_c] = values_partners_govs[key] + (cum_partner_gov[key_c] if key_c in cum_partner_gov else 0)

        if report_type == 'live':
            indicator.cumulative_values_live = {
                'months': cum_month,
                'partners': cum_partner,
                'govs': cum_gov,
                'partners_govs': cum_partner_gov
            }
        else:
            indicator.cumulative_values = {
                'months': cum_month,
                'partners': cum_partner,
                'govs': cum_gov,
                'partners_govs': cum_partner_gov
            }
        indicator.save()

    return indicators.count()


def reset_hpm_indicators_values():
    from internos.activityinfo.models import Indicator
    current_year = datetime.datetime.now().year

    indicators = Indicator.objects.filter(activity__database__reporting_year__year=current_year ,hpm_indicator=True)
    for indicator in indicators:
        indicator.values_hpm = {}
        # indicator.values_tags = {}
        indicator.cumulative_values_hpm = {}
        indicator.save()

    return indicators.count()


def calculate_indicators_cumulative_hpm():
    from internos.activityinfo.models import Indicator

    indicators = Indicator.objects.filter(hpm_indicator=True)

    for indicator in indicators:
        cum_month = {}

        values = indicator.values_hpm

        for month in values:
            c_value = 0
            try:
                for c_month in range(1, int(month) + 1):
                    c_value += float(values[str(c_month)])
                    cum_month[str(month)] = c_value
            except Exception:
                continue

        indicator.cumulative_values_hpm = {
            'months': cum_month,
        }
        indicator.save()

    return indicators.count()


def calculate_indicators_tags_hpm(ai_db):
    from internos.activityinfo.models import Indicator

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id).filter(hpm_indicator=True)

    for indicator in indicators.iterator():
        m_value = indicator.cumulative_values['months']
        if isinstance(m_value, dict):
            m_value = 0
        sub_indicators = indicator.summation_sub_indicators.all().only(
            'name',
            'values',
            'values_gov',
            'values_partners',
            'values_partners_gov',
            'values_live',
            'values_hpm',
            'values_tags',
            'cumulative_values',
        )
        tag_boys = sub_indicators.filter(tag_gender__name='Boys')
        tag_girls = sub_indicators.filter(tag_gender__name='Girls')
        tag_male = sub_indicators.filter(tag_gender__name='Male')
        tag_female = sub_indicators.filter(tag_gender__name='Female')

        tag_male_value = 0
        for ind_tag in tag_male:
            c_value = 0
            if 'months' in ind_tag.cumulative_values:
                c_value = ind_tag.cumulative_values['months']

            if isinstance(c_value, dict):
                c_value = 0

            tag_male_value += float(c_value)

        try:
            indicator.values_tags['male'] = float(tag_male_value) * 100 / float(m_value)
        except Exception as ex:
            # print(ex.message)
            indicator.values_tags['male'] = 0

        tag_female_value = 0
        for ind_tag in tag_female:
            c_value = 0
            if 'months' in ind_tag.cumulative_values:
                c_value = ind_tag.cumulative_values['months']

            if isinstance(c_value, dict):
                c_value = 0

            tag_female_value += c_value

        try:
            indicator.values_tags['female'] = float(tag_female_value) * 100 / float(m_value)
        except Exception as ex:
            # print(ex.message)
            indicator.values_tags['female'] = 0

        tag_boys_value = 0
        for ind_tag in tag_boys:
            c_value = 0
            if 'months' in ind_tag.cumulative_values:
                c_value = ind_tag.cumulative_values['months']

            if isinstance(c_value, dict):
                c_value = 0

            tag_boys_value += c_value

        try:
            indicator.values_tags['boys'] = float(tag_boys_value) * 100 / float(m_value)
        except Exception as ex:
            # print(ex.message)
            indicator.values_tags['boys'] = 0

        tag_girls_value = 0
        for ind_tag in tag_girls:
            c_value = 0
            if 'months' in ind_tag.cumulative_values:
                c_value = ind_tag.cumulative_values['months']

            if isinstance(c_value, dict):
                c_value = 0

            tag_girls_value += c_value

        try:
            indicator.values_tags['girls'] = float(tag_girls_value) * 100 / float(m_value)
        except Exception as ex:
            # print(ex.message)
            indicator.values_tags['girls'] = 0

        indicator.save()

    return indicators.count()


def calculate_indicators_all_tags(ai_db):

    calculate_indicators_tags(ai_db,True)
    calculate_indicators_tags(ai_db, False)


def calculate_indicators_tags(ai_db,sub_master=False):
    from internos.activityinfo.models import Indicator, IndicatorTag , ActivityReport

    # indicators = Indicator.objects.filter(hpm_indicator=True)
    if sub_master:
        indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id).\
            filter(Q(master_indicator_sub=True))
    else:
        indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id). \
            filter(Q(master_indicator=True) | Q(hpm_indicator=True))

    report = ActivityReport.objects.filter(database_id= ai_db.ai_id)

    if ai_db.is_funded_by_unicef:
        report = report.filter(funded_by__contains='UNICEF')

    partners = report.values('partner_id').order_by('partner_id').distinct('partner_id')
    governorates = report.values('location_adminlevel_governorate_code').distinct()
    sections = report.values('reporting_section').distinct()

    tags_gender = IndicatorTag.objects.filter(type='gender').only('id', 'name')
    tags_age = IndicatorTag.objects.filter(type='age').only('id', 'name')
    tags_nationality = IndicatorTag.objects.filter(type='nationality').only('id', 'name')
    tags_disability = IndicatorTag.objects.filter(type='disability').only('id', 'name')
    tags_programme = IndicatorTag.objects.filter(type='programme').only('id', 'name')

    for indicator in indicators.iterator():
        m_value = 0
        sub_indicators = indicator.summation_sub_indicators.all()
        for sub_indicator in sub_indicators:
            if sub_indicator.master_indicator:
                continue
            else:
                sub_indicators = sub_indicators | sub_indicator.summation_sub_indicators.all()
        try:
            m_value = indicator.cumulative_values['months']

        except Exception:
            continue

        if isinstance(m_value, dict):
            m_value = 0

        indicator.values_tags = {}
        sub_indicators =sub_indicators.only(
            'values',
            'values_partners',
            'values_partners_gov',
            'values_gov',
            'values_sections',
            'values_sections_partners',
            'values_sections_gov',
            'values_tags',
            'cumulative_values',
            'is_cumulative',
            'highest_values'
        )
        # ----------------------------- Gender tags --------------------------------
        for tag in tags_gender.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_gender_id=tag.id)
            if tag_sub_indicators:

                # -----------------------------  tags calculations per month  --------------------------------
                months_list = {}
                for mon in range(1, 13):
                    mon_value = 0
                    for ind_tag in tag_sub_indicators:
                        if str(mon) in ind_tag.values:
                            mon_value += ind_tag.values[str(mon)]
                            months_list['{}--{}'.format(mon,tag.name)] = mon_value
                indicator.values_tags['months_' + tag.name] = months_list

                # -----------------------------  tags calculations per partner  --------------------------------
                partners_list = {}
                for mon in range(1, 13):
                    for par in partners:
                        par = par['partner_id']
                        par_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(mon,par)
                            if key in ind_tag.values_partners:
                                par_value += ind_tag.values_partners[key]
                                partners_list['{}--{}--{}'.format(mon,par,tag.name)] = par_value
                indicator.values_tags['partners_' + tag.name] = partners_list

                # -----------------------------  tags calculations per governorate  --------------------------------
                govs_list = {}
                for mon in range(1, 13):
                    for gov in governorates:
                        gov = gov['location_adminlevel_governorate_code']
                        gov_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(mon, gov)
                            if key in ind_tag.values_gov:
                                gov_value += ind_tag.values_gov[key]
                                govs_list['{}--{}--{}'.format(mon, gov, tag.name)] = gov_value

                indicator.values_tags['govs_' + tag.name] = govs_list

                # -----------------------------  tags calculations per partner per gov --------------------------------
                partner_gov_list = {}
                for mon in range(1, 13):
                    for par in partners:
                        par = par['partner_id']
                        for gov in governorates:
                            gov = gov['location_adminlevel_governorate_code']
                            par_gv_value = 0
                            for ind_tag in tag_sub_indicators:
                                key = '{}-{}-{}'.format(mon, gov,par)
                                if key in ind_tag.values_partners_gov:
                                    par_gv_value += ind_tag.values_partners_gov[key]
                                    partner_gov_list['{}--{}--{}--{}'.format(mon, par, gov, tag.name)] = par_gv_value
                indicator.values_tags['partners_govs_' + tag.name] = partner_gov_list


                # ----------------------------- tags cumulative calculations per gov per partner   ---------------------
                cum_partner_gov = {}

                for gov in governorates:
                    gov = gov['location_adminlevel_governorate_code']
                    for par in partners:
                        par = par['partner_id']
                        par_gv_cum_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(gov, par)
                            if indicator.is_cumulative:
                                if 'partners_govs' in ind_tag.cumulative_values:
                                    if key in ind_tag.cumulative_values['partners_govs']:
                                        par_gv_cum_value += ind_tag.cumulative_values['partners_govs'][key]
                                        cum_partner_gov['{}--{}--{}'.format(par, gov, tag.name)] = par_gv_cum_value
                            else:
                                if 'partners_govs' in ind_tag.highest_values:
                                    if key in ind_tag.highest_values['partners_govs']:
                                        par_gv_cum_value += ind_tag.highest_values['partners_govs'][key]
                                        cum_partner_gov['{}--{}--{}'.format(par, gov, tag.name)] = par_gv_cum_value

                indicator.values_tags['cum_partner_gov_' + tag.name] = cum_partner_gov

                # -------------- tags cumulative calculations per partner ----------
                cum_partner = {}
                for par in partners:
                    par = par['partner_id']
                    par_cum_value = 0
                    for ind_tag in tag_sub_indicators:
                        key = '{}'.format(par)
                        if indicator.is_cumulative:
                            if 'partners' in ind_tag.cumulative_values:
                                if key in ind_tag.cumulative_values['partners']:
                                    par_cum_value += ind_tag.cumulative_values['partners'][key]
                                    cum_partner['{}--{}'.format(par, tag.name)] = par_cum_value
                        else:
                            if 'partners' in ind_tag.highest_values:
                                if key in ind_tag.highest_values['partners']:
                                    par_cum_value += ind_tag.highest_values['partners'][key]
                                    cum_partner['{}--{}'.format(par, tag.name)] = par_cum_value
                indicator.values_tags['cum_partners_' + tag.name] = cum_partner

                # -------------- tags cumulative calculations per gov ----------
                cum_gov = {}
                for gv in governorates:
                    gv = gv['location_adminlevel_governorate_code']
                    gov_cum_value = 0
                    for ind_tag in tag_sub_indicators:
                        key = '{}'.format(gv)
                        if indicator.is_cumulative:
                            if 'govs' in ind_tag.cumulative_values:
                                if key in ind_tag.cumulative_values['govs']:
                                    gov_cum_value += ind_tag.cumulative_values['govs'][key]
                                    cum_gov['{}--{}'.format(gv, tag.name)] = gov_cum_value
                        else:
                            if 'govs' in ind_tag.highest_values:
                                if key in ind_tag.highest_values['govs']:
                                    gov_cum_value += ind_tag.highest_values['govs'][key]
                                    cum_gov['{}--{}'.format(gv, tag.name)] = gov_cum_value
                indicator.values_tags['cum_govs_' + tag.name] = cum_gov

                indicator.save()

                # -----------------------------  tags calculations General value and percentage  -----------------------
                value = 0
                for ind_tag in tag_sub_indicators:
                    c_value = 0
                    if indicator.is_cumulative:
                        if 'months' in ind_tag.cumulative_values:
                            c_value = ind_tag.cumulative_values['months']
                    else:
                        if 'months' in ind_tag.highest_values:
                            c_value = ind_tag.highest_values['months']

                    if isinstance(c_value, dict):
                        c_value = 0

                    value += float(c_value)

                tag_name_per = '{}_per'.format(tag.name)
                try:
                    indicator.values_tags[tag_name_per] = float(value) * 100 / float(m_value)
                    indicator.values_tags[tag.name] = value
                except Exception as ex:
                    # print(ex.message)
                    indicator.values_tags[tag_name_per] = 0
                    indicator.values_tags[tag.name] = 0

        # ----------------------------- Age groups tags --------------------------------
        for tag in tags_age.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_age_id=tag.id)
            if tag_sub_indicators:

                partners_list = {}
                for mon in range(1, 13):
                    for par in partners:
                        par = par['partner_id']
                        par_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(mon, par)
                            if key in ind_tag.values_partners:
                                par_value += ind_tag.values_partners[key]
                                partners_list['{}--{}--{}'.format(mon, par, tag.name)] = par_value

                indicator.values_tags['partners_' + tag.name] = partners_list

                govs_list = {}
                for mon in range(1, 13):
                    for gov in governorates:
                        gov = gov['location_adminlevel_governorate_code']
                        gov_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(mon, gov)
                            if key in ind_tag.values_gov:
                                gov_value += ind_tag.values_gov[key]
                                govs_list['{}--{}--{}'.format(mon, gov, tag.name)] = gov_value

                indicator.values_tags['govs_' + tag.name] = govs_list

                partner_gov_list = {}
                for mon in range(1, 13):
                    for par in partners:
                        par = par['partner_id']
                        for gv in governorates:
                            gv = gv['location_adminlevel_governorate_code']
                            par_gv_value = 0
                            for ind_tag in tag_sub_indicators:
                                key = '{}-{}-{}'.format(mon, gv, par)
                                if key in ind_tag.values_partners_gov:
                                    par_gv_value += ind_tag.values_partners_gov[key]
                                    partner_gov_list['{}--{}--{}--{}'.format(mon, par, gv, tag.name)] = par_gv_value

                indicator.values_tags['partners_govs_' + tag.name] = partner_gov_list

                months_list = {}
                for mon in range(1, 13):
                    mon_value = 0
                    for ind_tag in tag_sub_indicators:
                        if str(mon) in ind_tag.values:
                            mon_value += ind_tag.values[str(mon)]
                    months_list['{}-{}'.format(mon, tag.name)] = mon_value

                indicator.values_tags['months_' + tag.name] = months_list

                cum_partner_gov = {}

                for gov in governorates:
                    gov = gov['location_adminlevel_governorate_code']
                    for par in partners:
                        par = par['partner_id']
                        par_gv_cum_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(gov, par)
                            if indicator.is_cumulative:
                                if 'partners_govs' in ind_tag.cumulative_values:
                                    if key in ind_tag.cumulative_values['partners_govs']:
                                        par_gv_cum_value += ind_tag.cumulative_values['partners_govs'][key]
                                        cum_partner_gov['{}--{}--{}'.format(par, gov, tag.name)] = par_gv_cum_value
                            else:
                                if 'partners_govs' in ind_tag.highest_values:
                                    if key in ind_tag.highest_values['partners_govs']:
                                        par_gv_cum_value += ind_tag.highest_values['partners_govs'][key]
                                        cum_partner_gov['{}--{}--{}'.format(par, gov, tag.name)] = par_gv_cum_value
                indicator.values_tags['cum_partner_gov_' + tag.name] = cum_partner_gov

                # -------------- tags cumulative calculations per partner ----------
                cum_partner = {}
                for par in partners:
                    par = par['partner_id']
                    par_cum_value = 0
                    for ind_tag in tag_sub_indicators:
                        key = '{}'.format(par)
                        if indicator.is_cumulative:
                            if 'partners' in ind_tag.cumulative_values:
                                if key in ind_tag.cumulative_values['partners']:
                                    par_cum_value += ind_tag.cumulative_values['partners'][key]
                                    cum_partner['{}--{}'.format(par, tag.name)] = par_cum_value
                        else:
                            if 'partners' in ind_tag.highest_values:
                                if key in ind_tag.highest_values['partners']:
                                    par_cum_value += ind_tag.highest_values['partners'][key]
                                    cum_partner['{}--{}'.format(par, tag.name)] = par_cum_value
                indicator.values_tags['cum_partners_' + tag.name] = cum_partner

                # -------------- tags cumulative calculations per gov ----------
                cum_gov = {}
                for gv in governorates:
                    gv = gv['location_adminlevel_governorate_code']
                    gov_cum_value = 0
                    for ind_tag in tag_sub_indicators:
                        key = '{}'.format(gv)
                        if indicator.is_cumulative:
                            if 'govs' in ind_tag.cumulative_values:
                                if key in ind_tag.cumulative_values['govs']:
                                    gov_cum_value += ind_tag.cumulative_values['govs'][key]
                                    cum_gov['{}--{}'.format(gv, tag.name)] = gov_cum_value
                        else:
                            if 'govs' in ind_tag.highest_values:
                                if key in ind_tag.highest_values['govs']:
                                    gov_cum_value += ind_tag.highest_values['govs'][key]
                                    cum_gov['{}--{}'.format(gv, tag.name)] = gov_cum_value
                indicator.values_tags['cum_govs_' + tag.name] = cum_gov
                indicator.save()

                value = 0
                for ind_tag in tag_sub_indicators:
                    c_value = 0
                    if indicator.is_cumulative:
                        if 'months' in ind_tag.cumulative_values:
                            c_value = ind_tag.cumulative_values['months']
                    else:
                        if 'months' in ind_tag.highest_values:
                            c_value = ind_tag.highest_values['months']

                    if isinstance(c_value, dict):
                        c_value = 0

                    value += float(c_value)

                tag_name_per = '{}_per'.format(tag.name)
                try:
                    indicator.values_tags[tag_name_per] = float(value) * 100 / float(m_value)
                    indicator.values_tags[tag.name] = value
                except Exception as ex:
                    # print(ex.message)
                    indicator.values_tags[tag_name_per] = 0
                    indicator.values_tags[tag.name] = 0

        # ----------------------------- Nationality tags --------------------------------
        for tag in tags_nationality.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_nationality_id=tag.id)
            if tag_sub_indicators:

                partners_list = {}
                for mon in range(1, 13):
                    for par in partners:
                        par = par['partner_id']
                        par_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(mon, par)
                            if key in ind_tag.values_partners:
                                par_value += ind_tag.values_partners[key]
                                partners_list['{}--{}--{}'.format(mon, par, tag.name)] = par_value

                indicator.values_tags['partners_' + tag.name] = partners_list

                govs_list = {}
                for mon in range(1, 13):
                    for gov in governorates:
                        gov = gov['location_adminlevel_governorate_code']
                        gov_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(mon, gov)
                            if key in ind_tag.values_gov:
                                gov_value += ind_tag.values_gov[key]
                                govs_list['{}--{}--{}'.format(mon, gov, tag.name)] = gov_value

                indicator.values_tags['govs_' + tag.name] = govs_list

                partner_gov_list = {}
                for mon in range(1, 13):
                    for par in partners:
                        par = par['partner_id']
                        for gv in governorates:
                            gv = gv['location_adminlevel_governorate_code']
                            par_gv_value = 0
                            for ind_tag in tag_sub_indicators:
                                key = '{}-{}-{}'.format(mon, gv, par)
                                if key in ind_tag.values_partners_gov:
                                    par_gv_value += ind_tag.values_partners_gov[key]
                                    partner_gov_list['{}--{}--{}--{}'.format(mon, par, gv, tag.name)] = par_gv_value

                indicator.values_tags['partners_govs_' + tag.name] = partner_gov_list

                months_list = {}
                for mon in range(1, 13):
                    mon_value = 0
                    for ind_tag in tag_sub_indicators:
                        if str(mon) in ind_tag.values:
                            mon_value += ind_tag.values[str(mon)]
                    months_list['{}-{}'.format(mon, tag.name)] = mon_value

                indicator.values_tags['months_' + tag.name] = months_list

                cum_partner_gov = {}

                for gov in governorates:
                    gov = gov['location_adminlevel_governorate_code']
                    for par in partners:
                        par = par['partner_id']
                        par_gv_cum_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(gov, par)
                            if indicator.is_cumulative:
                               if 'partners_govs' in ind_tag.cumulative_values:
                                    if key in ind_tag.cumulative_values['partners_govs']:
                                        par_gv_cum_value += ind_tag.cumulative_values['partners_govs'][key]
                                        cum_partner_gov['{}--{}--{}'.format(par, gov, tag.name)] = par_gv_cum_value
                            else:
                                if 'partners_govs' in ind_tag.highest_values:
                                    if key in ind_tag.highest_values['partners_govs']:
                                        par_gv_cum_value += ind_tag.highest_values['partners_govs'][key]
                                        cum_partner_gov['{}--{}--{}'.format(par, gov, tag.name)] = par_gv_cum_value
                indicator.values_tags['cum_partner_gov_' + tag.name] = cum_partner_gov


                # -------------- tags cumulative calculations per partner ----------
                cum_partner = {}
                for par in partners:
                    par = par['partner_id']
                    par_cum_value = 0
                    for ind_tag in tag_sub_indicators:
                        key = '{}'.format(par)
                        if indicator.is_cumulative:
                            if 'partners' in ind_tag.cumulative_values:
                                if key in ind_tag.cumulative_values['partners']:
                                    par_cum_value += ind_tag.cumulative_values['partners'][key]
                                    cum_partner['{}--{}'.format(par, tag.name)] = par_cum_value
                        else:
                            if 'partners' in ind_tag.highest_values:
                                if key in ind_tag.highest_values['partners']:
                                    par_cum_value += ind_tag.highest_values['partners'][key]
                                    cum_partner['{}--{}'.format(par, tag.name)] = par_cum_value
                indicator.values_tags['cum_partners_' + tag.name] = cum_partner

                # -------------- tags cumulative calculations per gov ----------
                cum_gov = {}
                for gv in governorates:
                    gv = gv['location_adminlevel_governorate_code']
                    gov_cum_value = 0
                    for ind_tag in tag_sub_indicators:
                        key = '{}'.format(gv)
                        if indicator.is_cumulative:
                            if 'govs' in ind_tag.cumulative_values:
                                if key in ind_tag.cumulative_values['govs']:
                                    gov_cum_value += ind_tag.cumulative_values['govs'][key]
                                    cum_gov['{}--{}'.format(gv, tag.name)] = gov_cum_value
                        else:
                            if 'govs' in ind_tag.highest_values:
                                if key in ind_tag.highest_values['govs']:
                                    gov_cum_value += ind_tag.highest_values['govs'][key]
                                    cum_gov['{}--{}'.format(gv, tag.name)] = gov_cum_value
                indicator.values_tags['cum_govs_' + tag.name] = cum_gov

                indicator.save()

                value = 0
                for ind_tag in tag_sub_indicators:
                    c_value = 0
                    if indicator.is_cumulative:
                        if 'months' in ind_tag.cumulative_values:
                            c_value = ind_tag.cumulative_values['months']
                    else:
                        if 'months' in ind_tag.highest_values:
                            c_value = ind_tag.highest_values['months']

                    if isinstance(c_value, dict):
                        c_value = 0

                    value += float(c_value)

                tag_name_per = '{}_per'.format(tag.name)
                try:
                    indicator.values_tags[tag_name_per] = float(value) * 100 / float(m_value)
                    indicator.values_tags[tag.name] = value
                except Exception as ex:
                    # print(ex.message)
                    indicator.values_tags[tag_name_per] = 0
                    indicator.values_tags[tag.name] = 0

        # ----------------------------- Programme tags --------------------------------
        for tag in tags_programme.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_programme_id=tag.id)
            if tag_sub_indicators:

                partners_list = {}
                for mon in range(1, 13):
                    for par in partners:
                        par = par['partner_id']
                        par_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(mon, par)
                            if key in ind_tag.values_partners:
                                par_value += ind_tag.values_partners[key]
                                partners_list['{}--{}--{}'.format(mon, par, tag.name)] = par_value

                indicator.values_tags['partners_' + tag.name] = partners_list

                govs_list = {}
                for mon in range(1, 13):
                    for gov in governorates:
                        gov = gov['location_adminlevel_governorate_code']
                        gov_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(mon, gov)
                            if key in ind_tag.values_gov:
                                gov_value += ind_tag.values_gov[key]
                                govs_list['{}--{}--{}'.format(mon, gov, tag.name)] = gov_value
                indicator.values_tags['govs_' + tag.name] = govs_list

                partner_gov_list = {}
                for mon in range(1, 13):
                    for par in partners:
                        par = par['partner_id']
                        for gv in governorates:
                            gv = gv['location_adminlevel_governorate_code']
                            par_gv_value = 0
                            for ind_tag in tag_sub_indicators:
                                key = '{}-{}-{}'.format(mon, gv, par)
                                if key in ind_tag.values_partners_gov:
                                    par_gv_value += ind_tag.values_partners_gov[key]
                                    partner_gov_list['{}--{}--{}--{}'.format(mon, par, gv, tag.name)] = par_gv_value

                indicator.values_tags['partners_govs_' + tag.name] = partner_gov_list

                months_list = {}
                for mon in range(1, 13):
                    mon_value = 0
                    for ind_tag in tag_sub_indicators:
                        if str(mon) in ind_tag.values:
                            mon_value += ind_tag.values[str(mon)]
                    months_list['{}-{}'.format(mon, tag.name)] = mon_value
                indicator.values_tags['months_' + tag.name] = months_list

                cum_partner_gov = {}

                for gov in governorates:
                    gov = gov['location_adminlevel_governorate_code']
                    for par in partners:
                        par = par['partner_id']
                        par_gv_cum_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(gov, par)
                            if indicator.is_cumulative:
                                if 'partners_govs' in ind_tag.cumulative_values:
                                    if key in ind_tag.cumulative_values['partners_govs']:
                                        par_gv_cum_value += ind_tag.cumulative_values['partners_govs'][key]
                                        cum_partner_gov['{}--{}--{}'.format(par, gov, tag.name)] = par_gv_cum_value
                            else:
                                if 'partners_govs' in ind_tag.highest_values:
                                    if key in ind_tag.highest_values['partners_govs']:
                                        par_gv_cum_value += ind_tag.highest_values['partners_govs'][key]
                                        cum_partner_gov['{}--{}--{}'.format(par, gov, tag.name)] = par_gv_cum_value
                indicator.values_tags['cum_partner_gov_' + tag.name] = cum_partner_gov

                cum_partner = {}
                for par in partners:
                    par = par['partner_id']
                    par_cum_value = 0
                    for ind_tag in tag_sub_indicators:
                        key = '{}'.format(par)
                        if indicator.is_cumulative:
                            if 'partners' in ind_tag.cumulative_values:
                                if key in ind_tag.cumulative_values['partners']:
                                    par_cum_value += ind_tag.cumulative_values['partners'][key]
                                    cum_partner['{}--{}'.format(par, tag.name)] = par_cum_value
                        else:
                            if 'partners' in ind_tag.highest_values:
                                if key in ind_tag.highest_values['partners']:
                                    par_cum_value += ind_tag.highest_values['partners'][key]
                                    cum_partner['{}--{}'.format(par, tag.name)] = par_cum_value
                indicator.values_tags['cum_partners_' + tag.name] = cum_partner

                # -------------- tags cumulative calculations per gov ----------
                cum_gov = {}
                for gv in governorates:
                    gv = gv['location_adminlevel_governorate_code']
                    gov_cum_value = 0
                    for ind_tag in tag_sub_indicators:
                        key = '{}'.format(gv)
                        if indicator.is_cumulative:
                            if 'govs' in ind_tag.cumulative_values:
                                if key in ind_tag.cumulative_values['govs']:
                                    gov_cum_value += ind_tag.cumulative_values['govs'][key]
                                    cum_gov['{}--{}'.format(gv, tag.name)] = gov_cum_value
                        else:
                            if 'govs' in ind_tag.highest_values:
                                if key in ind_tag.highest_values['govs']:
                                    gov_cum_value += ind_tag.highest_values['govs'][key]
                                    cum_gov['{}--{}'.format(gv, tag.name)] = gov_cum_value
                indicator.values_tags['cum_govs_' + tag.name] = cum_gov
                indicator.save()

                value = 0
                for ind_tag in tag_sub_indicators:
                    c_value = 0
                    if indicator.is_cumulative:
                        if 'months' in ind_tag.cumulative_values:
                            c_value = ind_tag.cumulative_values['months']
                    else:
                        if 'months' in ind_tag.highest_values:
                            c_value = ind_tag.highest_values['months']

                    if isinstance(c_value, dict):
                        c_value = 0

                    value += float(c_value)

                tag_name_per = '{}_per'.format(tag.name)
                try:
                    indicator.values_tags[tag_name_per] = float(value) * 100 / float(m_value)
                    indicator.values_tags[tag.name] = value
                except Exception as ex:
                    # print(ex.message)
                    indicator.values_tags[tag_name_per] = 0
                    indicator.values_tags[tag.name] = 0

        # ----------------------------- Disability tags --------------------------------

        for tag in tags_disability.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_disability_id=tag.id)

            if tag_sub_indicators:

                partners_list = {}
                for mon in range(1, 13):
                    for par in partners:
                        par = par['partner_id']
                        par_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(mon, par)
                            if key in ind_tag.values_partners:
                                par_value += ind_tag.values_partners[key]
                                partners_list['{}--{}--{}'.format(mon, par, tag.name)] = par_value
                indicator.values_tags['partners_' + tag.name] = partners_list

                govs_list = {}
                for mon in range(1, 13):
                    for gov in governorates:
                        gov = gov['location_adminlevel_governorate_code']
                        gov_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(mon, gov)
                            if key in ind_tag.values_gov:
                                gov_value += ind_tag.values_gov[key]
                                govs_list['{}--{}--{}'.format(mon, gov, tag.name)] = gov_value
                indicator.values_tags['govs_' + tag.name] = govs_list

                partner_gov_list = {}
                for mon in range(1, 13):
                    for par in partners:
                        par = par['partner_id']
                        for gv in governorates:
                            gv = gv['location_adminlevel_governorate_code']
                            par_gv_value = 0
                            for ind_tag in tag_sub_indicators:
                                key = '{}-{}-{}'.format(mon, gv, par)
                                if key in ind_tag.values_partners_gov:
                                    par_gv_value += ind_tag.values_partners_gov[key]
                                    partner_gov_list['{}--{}--{}--{}'.format(mon, par, gv, tag.name)] = par_gv_value

                indicator.values_tags['partners_govs_' + tag.name] = partner_gov_list

                months_list = {}
                for mon in range(1, 13):
                    mon_value = 0
                    for ind_tag in tag_sub_indicators:
                        if str(mon) in ind_tag.values:
                            mon_value += ind_tag.values[str(mon)]
                    months_list['{}-{}'.format(mon, tag.name)] = mon_value

                indicator.values_tags['months_' + tag.name] = months_list

                cum_partner_gov = {}
                for gov in governorates:
                    gov = gov['location_adminlevel_governorate_code']
                    for par in partners:
                        par = par['partner_id']
                        par_gv_cum_value = 0
                        for ind_tag in tag_sub_indicators:
                            key = '{}-{}'.format(gov, par)
                            if indicator.is_cumulative:
                                if 'partners_govs' in ind_tag.cumulative_values:
                                    if key in ind_tag.cumulative_values['partners_govs']:
                                        par_gv_cum_value += ind_tag.cumulative_values['partners_govs'][key]
                                        cum_partner_gov['{}--{}--{}'.format(par, gov, tag.name)] = par_gv_cum_value
                            else:
                                if 'partners_govs' in ind_tag.highest_values:
                                    if key in ind_tag.highest_values['partners_govs']:
                                        par_gv_cum_value += ind_tag.highest_values['partners_govs'][key]
                                        cum_partner_gov['{}--{}--{}'.format(par, gov, tag.name)] = par_gv_cum_value
                indicator.values_tags['cum_partner_gov_' + tag.name] = cum_partner_gov

                # -------------- tags cumulative calculations per partner ----------
                cum_partner = {}
                for par in partners:
                    par = par['partner_id']
                    par_cum_value = 0
                    for ind_tag in tag_sub_indicators:
                        key = '{}'.format(par)
                        if indicator.is_cumulative:
                            if 'partners' in ind_tag.cumulative_values:
                                if key in ind_tag.cumulative_values['partners']:
                                    par_cum_value += ind_tag.cumulative_values['partners'][key]
                                    cum_partner['{}--{}'.format(par, tag.name)] = par_cum_value
                        else:
                            if 'partners' in ind_tag.highest_values:
                                if key in ind_tag.highest_values['partners']:
                                    par_cum_value += ind_tag.highest_values['partners'][key]
                                    cum_partner['{}--{}'.format(par, tag.name)] = par_cum_value
                indicator.values_tags['cum_partners_' + tag.name] = cum_partner

                # -------------- tags cumulative calculations per gov ----------
                cum_gov = {}
                for gv in governorates:
                    gv = gv['location_adminlevel_governorate_code']
                    gov_cum_value = 0
                    for ind_tag in tag_sub_indicators:
                        key = '{}'.format(gv)
                        if indicator.is_cumulative:
                            if 'govs' in ind_tag.cumulative_values:
                                if key in ind_tag.cumulative_values['govs']:
                                    gov_cum_value += ind_tag.cumulative_values['govs'][key]
                                    cum_gov['{}--{}'.format(gv, tag.name)] = gov_cum_value
                        else:
                            if 'govs' in ind_tag.highest_values:
                                if key in ind_tag.highest_values['govs']:
                                    gov_cum_value += ind_tag.highest_values['govs'][key]
                                    cum_gov['{}--{}'.format(gv, tag.name)] = gov_cum_value
                indicator.values_tags['cum_govs_' + tag.name] = cum_gov
                indicator.save()

                value = 0
                for ind_tag in tag_sub_indicators:
                    c_value = 0
                    if indicator.is_cumulative:
                        if 'months' in ind_tag.cumulative_values:
                            c_value = ind_tag.cumulative_values['months']
                    else:
                        if 'months' in ind_tag.highest_values:
                            c_value = ind_tag.highest_values['months']

                    if isinstance(c_value, dict):
                        c_value = 0

                    value += float(c_value)

                tag_name_per = '{}_per'.format(tag.name)
                try:
                    indicator.values_tags[tag_name_per] = float(value) * 100 / float(m_value)
                    indicator.values_tags[tag.name] = value
                except Exception as ex:
                    # print(ex.message)
                    indicator.values_tags[tag_name_per] = 0
                    indicator.values_tags[tag.name] = 0

            indicator.save()

    return indicators.count()


def calculate_indicators_monthly_tags(ai_db):
    from internos.activityinfo.models import Indicator, IndicatorTag

    # indicators = Indicator.objects.filter(hpm_indicator=True)
    indicators = Indicator.objects.filter(Q(master_indicator=True) | Q(hpm_indicator=True))
    tags_gender = IndicatorTag.objects.filter(type='gender').only('id', 'name')
    tags_age = IndicatorTag.objects.filter(type='age').only('id', 'name')
    tags_nationality = IndicatorTag.objects.filter(type='nationality').only('id', 'name')
    tags_disability = IndicatorTag.objects.filter(type='disability').only('id', 'name')
    tags_programme = IndicatorTag.objects.filter(type='programme').only('id', 'name')

    today = datetime.date.today()
    first = today.replace(day=1)
    last_month = first - datetime.timedelta(days=1)
    month = int(last_month.strftime("%m"))
    # month = int(last_month.strftime("%m")) - 1
    month_str = str(month)

    month = 12
    month_str = 'December'

    for indicator in indicators.iterator():
        sub_indicators = indicator.summation_sub_indicators.all().only(
            'cumulative_values_hpm',
            'values',
        )

        if not indicator.cumulative_values_hpm:
            indicator.cumulative_values_hpm = {}
        if 'tags' not in indicator.cumulative_values_hpm:
            indicator.cumulative_values_hpm['tags'] = {}

        for tag in tags_gender.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_gender_id=tag.id)

            value = 0
            for ind_tag in tag_sub_indicators:
                c_value = 0
                if month_str in ind_tag.values:
                    c_value = ind_tag.values[month_str]

                value += float(c_value)

            key = '{}-{}'.format(month, tag.name)
            try:
                indicator.cumulative_values_hpm['tags'][key] = value
            except Exception as ex:
                # print(ex.message)
                indicator.cumulative_values_hpm['tags'][key] = 0

        for tag in tags_age.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_age_id=tag.id)

            value = 0
            for ind_tag in tag_sub_indicators:
                c_value = 0
                if month_str in ind_tag.values:
                    c_value = ind_tag.values[month_str]

                value += float(c_value)

            key = '{}-{}'.format(month_str, tag.name)
            try:
                indicator.cumulative_values_hpm['tags'][key] = value
            except Exception as ex:
                # print(ex.message)
                indicator.cumulative_values_hpm['tags'][key] = 0

        for tag in tags_nationality.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_nationality_id=tag.id)

            value = 0
            for ind_tag in tag_sub_indicators:
                c_value = 0
                if month_str in ind_tag.values:
                    c_value = ind_tag.values[month_str]

                value += float(c_value)

            key = '{}-{}'.format(month, tag.name)
            try:
                indicator.cumulative_values_hpm['tags'][key] = value
            except Exception as ex:
                # print(ex.message)
                indicator.cumulative_values_hpm['tags'][key] = 0

        for tag in tags_programme.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_programme_id=tag.id)

            value = 0
            for ind_tag in tag_sub_indicators:
                c_value = 0
                if month_str in ind_tag.values:
                    c_value = ind_tag.values[month_str]

                value += float(c_value)

            key = '{}-{}'.format(month, tag.name)
            try:
                indicator.cumulative_values_hpm['tags'][key] = value
            except Exception as ex:
                # print(ex.message)
                indicator.cumulative_values_hpm['tags'][key] = 0

        for tag in tags_disability.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_disability_id=tag.id)

            value = 0
            for ind_tag in tag_sub_indicators:
                c_value = 0
                if month_str in ind_tag.values:
                    c_value = ind_tag.values[month_str]

                value += float(c_value)

            key = '{}-{}'.format(month, tag.name)
            try:
                indicator.cumulative_values_hpm['tags'][key] = value
            except Exception as ex:
                # print(ex.message)
                indicator.cumulative_values_hpm['tags'][key] = 0

        indicator.save()

    return indicators.count()


def calculate_master_indicators_values_1(ai_db, report_type=None, sub_indicators=False):
    from django.db import connection
    from internos.activityinfo.models import Indicator

    if sub_indicators:
        indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                              master_indicator_sub=True)
    else:
        indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                              master_indicator=True)

    indicators = indicators.only(
        'summation_sub_indicators',
        'values',
        'values_gov',
        'values_partners',
        'values_partners_gov',
        'values_live',
        'values_gov_live',
        'values_partners_live',
        'values_partners_gov_live',
        'values_hpm',
        'values_sections',
        'values_sections_partners',
        'values_sections_gov',
        'values_sections_partners_gov',
        'values_sections_live',
        'values_sections_partners_live',
        'values_sections_gov_live',
        'values_sections_partners_gov_live'

    )

    rows_data = {}
    cursor = connection.cursor()
    cursor.execute(
        "SELECT distinct a1.id, a1.ai_indicator, ai.ai_indicator, ai.id, ai.values, ai.values_live, "
        "ai.values_gov, ai.values_gov_live, ai.values_partners, ai.values_partners_live, ai.values_partners_gov, "
        "ai.values_partners_gov_live, ai.values_sections, ai.values_sections_live, ai.values_sections_partners, "
        "ai.values_sections_partners_live, ai.values_sections_gov, ai.values_sections_gov_live, ai.values_sections_partners_gov,  "
        "ai.values_sections_partners_gov_live, a1.master_indicator, a1.master_indicator_sub "
        "FROM public.activityinfo_indicator a1, public.activityinfo_activity aa, "
        "public.activityinfo_indicator_summation_sub_indicators ais, public.activityinfo_indicator ai "
        "WHERE ai.activity_id = aa.id AND a1.id = ais.from_indicator_id AND ais.to_indicator_id = ai.id "
        "AND aa.database_id = %s AND (a1.master_indicator = true or a1.master_indicator_sub = true)",
        [ai_db.id])

    rows = cursor.fetchall()
    for row in rows:
        if row[0] not in rows_data:
            rows_data[row[0]] = {}

        rows_data[row[0]][row[3]] = row

    for indicator in indicators.iterator():
        values_month = {}
        values_partners = {}
        values_gov = {}
        values_partners_gov = {}
        values_sections = {}
        values_sections_partners = {}
        values_sections_gov = {}
        values_sections_partners_gov = {}

        if indicator.id in rows_data:
            indicator_values = rows_data[indicator.id]
            for key, key_values in indicator_values.items():
                sub_indicator_values = indicator_values[key]

                if report_type == 'live':
                    values = sub_indicator_values[5]  # values_live
                    values1 = sub_indicator_values[7]  # values_gov_live
                    values2 = sub_indicator_values[9]  # values_partners_live
                    values3 = sub_indicator_values[11]  # values_partners_gov_live
                    values4 = sub_indicator_values[13]  #values_sections_live
                    values5 = sub_indicator_values[15] # values_sections_partners_live
                    values6 = sub_indicator_values[17] # values_sections_gov_live
                    values7 = sub_indicator_values[19] # values_sections_partners_gov_live

                else:
                    values = sub_indicator_values[4]  # values
                    values1 = sub_indicator_values[6]  # values_gov
                    values2 = sub_indicator_values[8]  # values_partners
                    values3 = sub_indicator_values[10]  # values_partners_gov
                    values4 = sub_indicator_values[12]  # values_sections
                    values5 = sub_indicator_values[14]  # values_sections_partners
                    values6 = sub_indicator_values[16] # values_sections_gov
                    values7 = sub_indicator_values[18]  # values_sections_partners_gov


                for key in values:
                    val = values[key]
                    if key in values_month:
                        val = values_month[key] + val
                    values_month[key] = val
                    # if str(key) == str(reporting_month):
                    #     indicator.values_hpm[reporting_month] = val

                for key in values1:
                    val = values1[key]
                    if key in values_gov:
                        val = values_gov[key] + val
                    values_gov[key] = val

                for key in values2:
                    val = values2[key]
                    if key in values_partners:
                        val = values_partners[key] + val
                    values_partners[key] = val

                for key in values3:
                    val = values3[key]
                    if key in values_partners_gov:
                        val = values_partners_gov[key] + val
                    values_partners_gov[key] = val

                for key in values4:
                    val = values4[key]
                    if key in values_sections:
                        val = values_sections[key] + val
                    values_sections[key] = val

                for key in values5:
                    val = values5[key]
                    if key in values_sections_partners:
                        val = values_sections_partners[key] + val
                    values_sections_partners[key] = val

                for key in values6:
                    val = values6[key]
                    if key in values_sections_gov:
                        val = values_sections_gov[key] + val
                    values_sections_gov[key] = val

                for key in values7:
                    val = values7[key]
                    if key in values_sections_partners_gov:
                        val = values_sections_partners_gov[key] + val
                    values_sections_partners_gov[key] = val

                if report_type == 'live':
                    indicator.values_live = values_month
                    indicator.values_gov_live = values_gov
                    indicator.values_partners_live = values_partners
                    indicator.values_partners_gov_live = values_partners_gov
                    indicator.values_sections_live= values_sections
                    indicator.values_sections_partners_live= values_sections_partners
                    indicator.values_sections_gov_live = values_sections_gov
                    indicator.values_sections_partners_gov_live = values_sections_partners_gov
                else:
                    indicator.values = values_month
                    indicator.values_gov = values_gov
                    indicator.values_partners = values_partners
                    indicator.values_partners_gov = values_partners_gov
                    indicator.values_sections = values_sections
                    indicator.values_sections_partners = values_sections_partners
                    indicator.values_sections_gov = values_sections_gov
                    indicator.values_sections_partners_gov = values_sections_partners_gov

            indicator.save()


#  todo to remove
def calculate_master_indicators_values(ai_db, report_type=None, sub_indicators=False):
    from internos.activityinfo.models import Indicator, ActivityReport, LiveActivityReport

    if sub_indicators:
        indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                              master_indicator_sub=True)
    else:
        indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                              master_indicator=True)

    indicators = indicators.only(
        'values',
        'values_gov',
        'values_partners',
        'values_partners_gov',
        'values_live',
        'values_gov_live',
        'values_partners_live',
        'values_partners_gov_live',
        'values_hpm',
    )

    # last_month = int(datetime.datetime.now().strftime("%m"))
    last_month = get_extraction_month(ai_db)
    # last_month = 13

    if report_type == 'live':
        report = LiveActivityReport.objects.filter(database_id=ai_db.ai_id)
        last_month = get_live_extraction_month(ai_db) + 1
        # last_month = last_month + 1
    else:
        report = ActivityReport.objects.filter(database_id=ai_db.ai_id)

    if ai_db.is_funded_by_unicef:
        report = report.filter(funded_by='UNICEF')

    report = report.only('partner_id', 'location_adminlevel_governorate_code')

    partners = report.values('partner_id').order_by('partner_id').distinct('partner_id')
    governorates = report.values('location_adminlevel_governorate_code').order_by('location_adminlevel_governorate_code').distinct('location_adminlevel_governorate_code')
    governorates1 = report.values('location_adminlevel_governorate_code').order_by('location_adminlevel_governorate_code').distinct('location_adminlevel_governorate_code')

    for indicator in indicators.iterator():
        for month in range(1, last_month):
            month = str(month)
            values_month = 0
            values_gov = {}
            values_partners = {}
            values_partners_gov = {}

            sub_indicators = indicator.summation_sub_indicators.all().only(
                'values',
                'values_gov',
                'values_partners',
                'values_partners_gov',
                'values_live',
                'values_gov_live',
                'values_partners_live',
                'values_partners_gov_live',
                'values_hpm',
            )
            for sub_ind in sub_indicators.iterator():
                if report_type == 'live':
                    values_month += float(sub_ind.values_live[month]) if month in sub_ind.values_live else 0
                else:
                    values_month += float(sub_ind.values[month]) if month in sub_ind.values else 0

                for gov1 in governorates1:
                    key = "{}-{}".format(month, gov1['location_adminlevel_governorate_code'])
                    if report_type == 'live':
                        value = float(sub_ind.values_gov_live[key]) if key in sub_ind.values_gov_live else 0
                    else:
                        value = float(sub_ind.values_gov[key]) if key in sub_ind.values_gov else 0
                    values_gov[key] = values_gov[key] + value if key in values_gov else value

                for partner in partners:
                    key1 = "{}-{}".format(month, partner['partner_id'])
                    if report_type == 'live':
                        value = float(sub_ind.values_partners_live[key1]) if key1 in sub_ind.values_partners_live else 0
                    else:
                        value = float(sub_ind.values_partners[key1]) if key1 in sub_ind.values_partners else 0
                    values_partners[key1] = values_partners[key1] + value if key1 in values_partners else value

                    for gov in governorates:
                        key2 = "{}-{}-{}".format(month, partner['partner_id'], gov['location_adminlevel_governorate_code'])
                        if report_type == 'live':
                            value = float(sub_ind.values_partners_gov_live[key2]) if key2 in sub_ind.values_partners_gov_live else 0
                        else:
                            value = float(sub_ind.values_partners_gov[key2]) if key2 in sub_ind.values_partners_gov else 0
                        values_partners_gov[key2] = values_partners_gov[key2] + value if key2 in values_partners_gov else value

            if report_type == 'live':
                indicator.values_live[month] = values_month
                indicator.values_gov_live.update(values_gov)
                indicator.values_partners_live.update(values_partners)
                indicator.values_partners_gov_live.update(values_partners_gov)
            else:
                # if month == reporting_month:
                #     indicator.values_hpm[reporting_month] = values_month
                indicator.values[month] = values_month
                indicator.values_gov.update(values_gov)
                indicator.values_partners.update(values_partners)
                indicator.values_partners_gov.update(values_partners_gov)

        indicator.save()


def calculate_indicators_values_percentage_1(ai_db, report_type=None):
    from django.db import connection
    from internos.activityinfo.models import Indicator

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                          calculated_indicator=True).only(
        'denominator_indicator',
        'numerator_indicator',
        'denominator_multiplication',
        'values',
        'values_gov',
        'values_partners',
        'values_partners_gov',
        'values_live',
        'values_gov_live',
        'values_partners_live',
        'values_partners_gov_live',
        'values_sections',
        'values_sections_partners',
        'values_sections_gov',
        'values_sections_partners_gov',
        'values_sections_live',
        'values_sections_partners_live',
        'values_sections_gov_live',
        'values_sections_partners_gov_live'
    )

    rows_data = {}
    cursor = connection.cursor()
    cursor.execute(
        "SELECT distinct a1.id, a1.calculated_percentage, aa.name, ai.id, ai.values, ai.values_live, "
        "ai.values_gov, ai.values_gov_live, ai.values_partners, ai.values_partners_live, ai.values_partners_gov, "
        "ai.values_partners_gov_live , ai.values_sections, ai.values_sections_live, ai.values_sections_partners, "
        "ai.values_sections_partners_live, ai.values_sections_gov, ai.values_sections_gov_live, ai.values_sections_partners_gov, "
        "ai.values_sections_partners_gov_live "
        "FROM public.activityinfo_indicator a1, public.activityinfo_activity aa, "
        "public.activityinfo_indicator_sub_indicators ais, public.activityinfo_indicator ai "
        "WHERE ai.activity_id = aa.id AND a1.id = ais.from_indicator_id AND ais.to_indicator_id = ai.id "
        "AND aa.database_id = %s AND a1.calculated_indicator = true",
        [ai_db.id])

    rows = cursor.fetchall()
    for row in rows:
        rows_data[row[0]] = row

    for indicator in indicators.iterator():
        values_month = {}
        values_partners = {}
        values_gov = {}
        values_partners_gov = {}
        values_sections = {}
        values_sections_partners = {}
        values_sections_gov = {}
        values_sections_partners_gov = {}

        if indicator.id in rows_data:
            indicator_values = rows_data[indicator.id]
            reporting_level = indicator_values[2]
            calculated_percentage = indicator_values[1]

            if report_type == 'live':
                values = indicator_values[5]  # values_live
                values1 = indicator_values[7]  # values_gov_live
                values2 = indicator_values[9]  # values_partners_live
                values3 = indicator_values[11]  # values_partners_gov_live
                values4 = indicator_values[13]  # values_sections_live
                values5 = indicator_values[15]  # values_sections_partners_live
                values6 = indicator_values[17]  # values_sections_gov_live
                values7 = indicator_values[19]  # values_sections_partners_gov_live
            else:
                values = indicator_values[4]  # values
                values1 = indicator_values[6]  # values_gov
                values2 = indicator_values[8]  # values_partners
                values3 = indicator_values[10]  # values_partners_gov
                values4 = indicator_values[12]  # values_sections
                values5 = indicator_values[14]  # values_sections_partners
                values6 = indicator_values[16]  # values_sections_gov
                values7 = indicator_values[18]  # values_sections_partners_gov

            for key in values:
                val = values[key]
                try:
                    if reporting_level == 'Municipality level':
                        values_month[key] = val * calculated_percentage / 100
                    elif reporting_level == 'Site level':
                        values_month[key] = val * calculated_percentage / 100
                except Exception as ex:
                    values_month[key] = 0

            for key in values1:
                val = values1[key]
                try:
                    if reporting_level == 'Municipality level':
                        values_gov[key] = val * calculated_percentage / 100
                    elif reporting_level == 'Site level':
                        values_gov[key] = val * calculated_percentage / 100
                except Exception as ex:
                    values_gov[key] = 0

            for key in values2:
                val = values2[key]
                try:
                    if reporting_level == 'Municipality level':
                        values_partners[key] = val * calculated_percentage / 100
                    elif reporting_level == 'Site level':
                        values_partners[key] = val * calculated_percentage / 100
                except Exception as ex:
                    values_partners[key] = 0

            for key in values3:
                val = values3[key]
                try:
                    if reporting_level == 'Municipality level':
                        values_partners_gov[key] = val * calculated_percentage / 100
                    elif reporting_level == 'Site level':
                        values_partners_gov[key] = val * calculated_percentage / 100
                except Exception as ex:
                    values_partners_gov[key] = 0

            for key in values4:
                val = values4[key]
                try:
                    if reporting_level == 'Municipality level':
                        values_sections[key] = val * calculated_percentage / 100
                    elif reporting_level == 'Site level':
                        values_sections[key] = val * calculated_percentage / 100
                except Exception as ex:
                    values_sections[key] = 0

            for key in values5:
                val = values5[key]
                try:
                    if reporting_level == 'Municipality level':
                        values_sections_partners[key] = val * calculated_percentage / 100
                    elif reporting_level == 'Site level':
                        values_sections_partners[key] = val * calculated_percentage / 100
                except Exception as ex:
                    values_sections_partners[key] = 0

            for key in values6:
                val = values6[key]
                try:
                    if reporting_level == 'Municipality level':
                        values_sections_gov[key] = val * calculated_percentage / 100
                    elif reporting_level == 'Site level':
                        values_sections_gov[key] = val * calculated_percentage / 100
                except Exception as ex:
                    values_sections_gov[key] = 0

            for key in values7:
                val = values7[key]
                try:
                    if reporting_level == 'Municipality level':
                        values_sections_partners_gov[key] = val * calculated_percentage / 100
                    elif reporting_level == 'Site level':
                        values_sections_partners_gov[key] = val * calculated_percentage / 100
                except Exception as ex:
                    values_sections_partners_gov[key] = 0

            if report_type == 'live':
                indicator.values_live = values_month
                indicator.values_gov_live = values_gov
                indicator.values_partners_live = values_partners
                indicator.values_partners_gov_live = values_partners_gov
                indicator.values_sections_live = values_sections
                indicator.values_sections_partners_live = values_sections_partners
                indicator.values_sections_gov_live = values_sections_gov
                indicator.values_sections_partners_gov_live = values_sections_partners_gov
            else:
                indicator.values = values_month
                indicator.values_gov = values_gov
                indicator.values_partners = values_partners
                indicator.values_partners_gov = values_partners_gov
                indicator.values_sections = values_sections
                indicator.values_sections_partners = values_sections_partners
                indicator.values_sections_gov = values_sections_gov
                indicator.values_sections_partners_gov = values_sections_partners_gov

            indicator.save()


def calculate_master_indicators_values_percentage(ai_db, report_type=None):
    from internos.activityinfo.models import Indicator, ActivityReport, LiveActivityReport

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                          master_indicator=True,
                                          measurement_type='percentage').only(
        'denominator_indicator',
        'numerator_indicator',
        'denominator_multiplication',
        'values',
        'values_gov',
        'values_partners',
        'values_partners_gov',
        'values_live',
        'values_gov_live',
        'values_partners_live',
        'values_partners_gov_live',
        'values_hpm',
        'values_sections',
        'values_sections_partners',
        'values_sections_gov',
        'values_sections_partners_gov',
        'values_sections_live',
        'values_sections_partners_live',
        'values_sections_gov_live',
        'values_sections_partners_gov_live'

    )
    last_month = get_extraction_month(ai_db)
    # last_month = int(datetime.datetime.now().strftime("%m"))

    if report_type == 'live':
        report = LiveActivityReport.objects.filter(database_id=ai_db.ai_id)
        # last_month = last_month + 1
        last_month = get_live_extraction_month(ai_db) + 1
    else:
        report = ActivityReport.objects.filter(database_id=ai_db.ai_id)
    if ai_db.is_funded_by_unicef:
        report = report.filter(funded_by='UNICEF')

    report = report.only('partner_id', 'location_adminlevel_governorate_code','reporting_section')


    partners = report.values('partner_id').order_by('partner_id').distinct('partner_id')
    governorates = report.values('location_adminlevel_governorate_code').order_by('location_adminlevel_governorate_code').distinct('location_adminlevel_governorate_code')
    governorates1 = report.values('location_adminlevel_governorate_code').order_by('location_adminlevel_governorate_code').distinct('location_adminlevel_governorate_code')
    sections = report.values('reporting_section').distinct()
    # last_month = 13

    for indicator in indicators.iterator():

        try:
            if report_type == 'live':
                denominator = denominator_indicator.cumulative_values_live['months'] if 'months' in denominator_indicator.cumulative_values_live['months'] else 0
                numerator = numerator_indicator.cumulative_values_live['months'] if 'months' in numerator_indicator.cumulative_values_live['months'] else 0
            else:
                denominator = denominator_indicator.cumulative_values['months'] if 'months' in denominator_indicator.cumulative_values['months'] else 0
                numerator = numerator_indicator.cumulative_values['months'] if 'months' in numerator_indicator.cumulative_values['months'] else 0
            cumulative_months = numerator / denominator
        except Exception as ex:
            logger.error(ex.message)
            cumulative_months = 0

        for month in range(1, last_month):
            month = str(month)
            values_gov = {}
            values_partners = {}
            values_partners_gov = {}
            values_sections={}
            values_sections_partners = {}
            values_sections_gov = {}
            values_sections_partners_gov = {}
            denominator_indicator = indicator.denominator_indicator
            numerator_indicator = indicator.numerator_indicator
            if not denominator_indicator or not numerator_indicator:
                continue
            try:
                if report_type == 'live':
                    denominator = denominator_indicator.values_live[month] if month in denominator_indicator.values_live else 0
                    numerator = numerator_indicator.values_live[month] if month in numerator_indicator.values_live else 0
                else:
                    denominator = denominator_indicator.values[month] if month in denominator_indicator.values else 0
                    numerator = numerator_indicator.values[month] if month in numerator_indicator.values else 0
                values_month = numerator / denominator
            except Exception as ex:
                logger.error(ex.message)
                values_month = 0

            for gov1 in governorates1:
                key = "{}-{}".format(month, gov1['location_adminlevel_governorate_code'])
                try:
                    if report_type == 'live':
                        denominator = denominator_indicator.values_gov_live[key] if key in denominator_indicator.values_gov_live else 0
                        numerator = numerator_indicator.values_gov_live[key] if key in numerator_indicator.values_gov_live else 0
                    else:
                        denominator = denominator_indicator.values_gov[key] if key in denominator_indicator.values_gov else 0
                        numerator = numerator_indicator.values_gov[key] if key in numerator_indicator.values_gov else 0
                    values_gov[key] = numerator / denominator
                except Exception as ex:
                    logger.error(ex.message)
                    values_gov[key] = 0

                # for section in sections:
                #     if 'reporting_section' in section and section['reporting_section']:
                #         key3 = "{}-{}-{}".format(month, section['reporting_section'],gov1['location_adminlevel_governorate_code'])
                #
                #         try:
                #             if report_type == 'live':
                #                 denominator = denominator_indicator.values_sections_gov_live[
                #                     key3] if key3 in denominator_indicator.values_sections_gov_live else 0
                #                 numerator = numerator_indicator.values_sections_gov_live[
                #                     key3] if key3 in numerator_indicator.values_sections_gov_live else 0
                #             else:
                #                 denominator = denominator_indicator.values_sections_gov[
                #                     key3] if key3 in denominator_indicator.values_sections_gov else 0
                #                 numerator = numerator_indicator.values_sections_gov[
                #                     key3] if key3 in numerator_indicator.values_sections_gov else 0
                #             values_sections_gov[key3] = numerator / denominator
                #         except Exception:
                #             values_sections_gov[key3] = 0

            # for section in sections:
            #     if 'reporting_section' in section and section['reporting_section']:
            #         key4 = "{}-{}".format(month, section['reporting_section'])
            #
            #         try:
            #             if report_type == 'live':
            #                 denominator = denominator_indicator.values_sections_live[
            #                     key4] if key4 in denominator_indicator.values_sections_live else 0
            #                 numerator = numerator_indicator.values_sections_live[
            #                     key4] if key4 in numerator_indicator.values_sections_live else 0
            #             else:
            #                 denominator = denominator_indicator.values_sections[
            #                     key4] if key4 in denominator_indicator.values_sections else 0
            #                 numerator = numerator_indicator.values_sections[
            #                     key4] if key4 in numerator_indicator.values_sections else 0
            #             values_sections[key4] = numerator / denominator
            #         except Exception:
            #             values_sections[key4] = 0

            for partner in partners:
                key1 = "{}-{}".format(month, partner['partner_id'])

                try:
                    if report_type == 'live':
                        denominator = denominator_indicator.values_partners_live[key1] if key1 in denominator_indicator.values_partners_live else 0
                        numerator = numerator_indicator.values_partners_live[key1] if key1 in numerator_indicator.values_partners_live else 0
                    else:
                        denominator = denominator_indicator.values_partners[key1] if key1 in denominator_indicator.values_partners else 0
                        numerator = numerator_indicator.values_partners[key1] if key1 in numerator_indicator.values_partners else 0
                    values_partners[key1] = numerator / denominator
                except Exception as ex:
                    logger.error(ex.message)
                    values_partners[key1] = 0

                for gov in governorates:
                    key2 = "{}-{}-{}".format(month, partner['partner_id'], gov['location_adminlevel_governorate_code'])
                    try:
                        if report_type == 'live':
                            denominator = denominator_indicator.values_partners_gov_live[key2] if key2 in denominator_indicator.values_partners_gov_live else 0
                            numerator = numerator_indicator.values_partners_gov_live[key2] if key2 in numerator_indicator.values_partners_gov_live else 0
                        else:
                            denominator = denominator_indicator.values_partners_gov[key2] if key2 in denominator_indicator.values_partners_gov else 0
                            numerator = numerator_indicator.values_partners_gov[key2] if key2 in numerator_indicator.values_partners_gov else 0
                        values_partners_gov[key2] = numerator / denominator
                    except Exception as ex:
                        logger.error(ex.message)
                        values_partners_gov[key2] = 0

                    # for section in sections:
                    #     key5 = "{}-{}-{}-{}".format(month, section['reporting_section'], partner['partner_id'],gov['location_adminlevel_governorate_code'])
                    #
                    #     try:
                    #         if report_type == 'live':
                    #             denominator = denominator_indicator.values_sections_partners_gov_live[
                    #                 key5] if key5 in denominator_indicator.values_sections_partners_gov_live else 0
                    #             numerator = numerator_indicator.values_sections_partners_gov_live[
                    #                 key5] if key5 in numerator_indicator.values_sections_partners_gov_live else 0
                    #         else:
                    #             denominator = denominator_indicator.values_sections_partners_gov[
                    #                 key5] if key5 in denominator_indicator.values_sections_partners_gov else 0
                    #             numerator = numerator_indicator.values_sections_partners_gov[
                    #                 key5] if key5 in numerator_indicator.values_sections_partners_gov else 0
                    #         values_sections_partners_gov[key5] = numerator / denominator
                    #     except Exception:
                    #         values_sections_partners_gov[key5] = 0

                # for section in sections:
                #     if 'reporting_section' in section and section['reporting_section']:
                #         key6 = "{}-{}-{}".format(month, section['reporting_section'],partner['partner_id'])
                #
                #         try:
                #             if report_type == 'live':
                #                 denominator = denominator_indicator.values_sections_partners_live[
                #                     key6] if key6 in denominator_indicator.values_sections_partners_live else 0
                #                 numerator = numerator_indicator.values_sections_partners_live[
                #                     key6] if key6 in numerator_indicator.values_sections_partners_live else 0
                #             else:
                #                 denominator = denominator_indicator.values_sections_partners[
                #                     key6] if key6 in denominator_indicator.values_sections_partners else 0
                #                 numerator = numerator_indicator.values_sections_partners[
                #                     key6] if key6 in numerator_indicator.values_sections_partners else 0
                #             values_sections_partners[key6] = numerator / denominator
                #         except Exception:
                #             values_sections_partners[key6] = 0

            if report_type == 'live':
                indicator.values_live[month] = values_month
                indicator.values_gov_live.update(values_gov)
                indicator.values_partners_live.update(values_partners)
                indicator.values_partners_gov_live.update(values_partners_gov)
                indicator.values_sections_live.update(values_sections)
                indicator.values_sections_partners_live.update(values_sections_partners)
                indicator.values_sections_gov_live.update(values_sections_gov)
                indicator.values_sections_partners_gov_live.update(values_sections_partners_gov)

            else:
                # if month == reporting_month:
                #     indicator.values_hpm[reporting_month] = values_month
                indicator.values[month] = values_month
                indicator.values_gov.update(values_gov)
                indicator.values_partners.update(values_partners)
                indicator.values_partners_gov.update(values_partners_gov)
                indicator.values_sections.update(values_sections)
                indicator.values_sections_partners.update(values_sections_partners)
                indicator.values_sections_gov.update(values_sections_gov)
                indicator.values_sections_partners_gov.update(values_sections_partners_gov)

        if report_type == 'live':
            indicator.cumulative_values_live['months'] = cumulative_months
        else:
            indicator.cumulative_values['months'] = cumulative_months

        if report_type == 'live':
            indicator.cumulative_values_live['months'] = cumulative_months
        else:
            indicator.cumulative_values['months'] = cumulative_months

        indicator.save()


def calculate_master_indicators_values_denominator_multiplication(ai_db, report_type=None):
    from internos.activityinfo.models import Indicator, ActivityReport, LiveActivityReport

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                          master_indicator=True,
                                          measurement_type='percentage_x').only(
        'denominator_indicator',
        'numerator_indicator',
        'denominator_multiplication',
        'values',
        'values_gov',
        'values_partners',
        'values_partners_gov',
        'values_live',
        'values_gov_live',
        'values_partners_live',
        'values_partners_gov_live',
        'values_hpm',
        'values_sections',
        'values_sections_partners',
        'values_sections_gov',
        'values_sections_partners_gov',
        'values_sections_live',
        'values_sections_partners_live',
        'values_sections_gov_live',
        'values_sections_partners_gov_live'
    )

    last_month = int(datetime.datetime.now().strftime("%m"))
    last_month = get_extraction_month(ai_db)

    if report_type == 'live':
        report = LiveActivityReport.objects.filter(database_id=ai_db.ai_id)
        # last_month = last_month + 1
        last_month = get_live_extraction_month(ai_db) + 1
    else:
        report = ActivityReport.objects.filter(database_id=ai_db.ai_id)
    if ai_db.is_funded_by_unicef:
        report = report.filter(funded_by='UNICEF')

    report = report.only('partner_id', 'location_adminlevel_governorate_code')

    partners = report.values('partner_id').order_by('partner_id').distinct('partner_id')
    governorates = report.values('location_adminlevel_governorate_code').order_by(
        'location_adminlevel_governorate_code').distinct('location_adminlevel_governorate_code')
    governorates1 = report.values('location_adminlevel_governorate_code').order_by(
        'location_adminlevel_governorate_code').distinct('location_adminlevel_governorate_code')

    sections = report.values('reporting_section').distinct()
    # last_month = 13

    for indicator in indicators.iterator():
        for month in range(1, last_month):
            month = str(month)
            values_gov = {}
            values_partners = {}
            values_partners_gov = {}
            values_sections = {}
            values_sections_partners = {}
            values_sections_gov = {}
            values_sections_partners_gov = {}

            denominator_indicator = indicator.denominator_indicator
            numerator_indicator = indicator.numerator_indicator
            denominator_multiplication = indicator.denominator_multiplication
            if not denominator_indicator or not numerator_indicator:
                continue
            try:
                if report_type == 'live':
                    denominator = denominator_indicator.values_live[month] if month in denominator_indicator.values_live else 0
                    numerator = numerator_indicator.values_live[month] if month in numerator_indicator.values_live else 0
                else:
                    denominator = denominator_indicator.values[month] if month in denominator_indicator.values else 0
                    numerator = numerator_indicator.values[month] if month in numerator_indicator.values else 0
                denominator = denominator * denominator_multiplication
                values_month = numerator / denominator
            except Exception as ex:
                values_month = 0

            for gov1 in governorates1:
                key = "{}-{}".format(month, gov1['location_adminlevel_governorate_code'])
                try:
                    if report_type == 'live':
                        denominator = denominator_indicator.values_gov_live[key] if key in denominator_indicator.values_gov_live else 0
                        numerator = numerator_indicator.values_gov_live[key] if key in numerator_indicator.values_gov_live else 0
                    else:
                        denominator = denominator_indicator.values_gov[key] if key in denominator_indicator.values_gov else 0
                        numerator = numerator_indicator.values_gov[key] if key in numerator_indicator.values_gov else 0
                    denominator = denominator * denominator_multiplication
                    values_gov[key] = numerator / denominator
                except Exception:
                    values_gov[key] = 0

            # for section in sections:
            #     if 'reporting_section' in section and section['reporting_section']:
            #         key = "{}-{}".format(month, section['reporting_section'])
            #         try:
            #             if report_type == 'live':
            #                 denominator = denominator_indicator.values_sections_live[
            #                     key] if key in denominator_indicator.values_sections_live else 0
            #                 numerator = numerator_indicator.values_sections_live[
            #                     key] if key in numerator_indicator.values_sections_live else 0
            #             else:
            #                 denominator = denominator_indicator.values_sections[
            #                     key] if key in denominator_indicator.values_sections else 0
            #                 numerator = numerator_indicator.values_sections[key] if key in numerator_indicator.values_sections else 0
            #             denominator = denominator * denominator_multiplication
            #             values_sections[key] = numerator / denominator
            #         except Exception:
            #             values_sections[key] = 0

            for partner in partners:
                key1 = "{}-{}".format(month, partner['partner_id'])

                try:
                    if report_type == 'live':
                        denominator = denominator_indicator.values_partners_live[key1] if key1 in denominator_indicator.values_partners_live else 0
                        numerator = numerator_indicator.values_partners_live[key1] if key1 in numerator_indicator.values_partners_live else 0
                    else:
                        denominator = denominator_indicator.values_partners[key1] if key1 in denominator_indicator.values_partners else 0
                        numerator = numerator_indicator.values_partners[key1] if key1 in numerator_indicator.values_partners else 0
                    denominator = denominator * denominator_multiplication
                    values_partners[key1] = numerator / denominator
                except Exception:
                    values_partners[key1] = 0

                for gov in governorates:
                    key2 = "{}-{}-{}".format(month, partner['partner_id'], gov['location_adminlevel_governorate_code'])
                    try:
                        if report_type == 'live':
                            denominator = denominator_indicator.values_partners_gov_live[key2] if key2 in denominator_indicator.values_partners_gov_live else 0
                            numerator = numerator_indicator.values_partners_gov_live[key2] if key2 in numerator_indicator.values_partners_gov_live else 0
                        else:
                            denominator = denominator_indicator.values_partners_gov[key2] if key2 in denominator_indicator.values_partners_gov else 0
                            numerator = numerator_indicator.values_partners_gov[key2] if key2 in numerator_indicator.values_partners_gov else 0
                        denominator = denominator * denominator_multiplication
                        values_partners_gov[key2] = numerator / denominator
                    except Exception:
                        values_partners_gov[key2] = 0

                    # for section in sections:
                    #     key6 = "{}-{}-{}".format(month, section['reporting_section'],
                    #                                 gov['location_adminlevel_governorate_code'])
                    #     try:
                    #         if report_type == 'live':
                    #             denominator = denominator_indicator.values_sections_gov_live[
                    #                 key6] if key6 in denominator_indicator.values_sections_gov_live else 0
                    #             numerator = numerator_indicator.values_sections_gov_live[
                    #                 key6] if key6 in numerator_indicator.values_sections_gov_live else 0
                    #         else:
                    #             denominator = denominator_indicator.values_sections_gov[
                    #                 key6] if key6 in denominator_indicator.values_sections_gov else 0
                    #             numerator = numerator_indicator.values_sections_gov[
                    #                 key6] if key6 in numerator_indicator.values_sections_gov else 0
                    #         denominator = denominator * denominator_multiplication
                    #         values_sections_gov[key6] = numerator / denominator
                    #     except Exception:
                    #         values_sections_gov[key6] = 0

                        # for section in sections:
                        #     key3 = "{}-{}-{}-{}".format(month, section['reporting_section'], partner['partner_id'],gov['location_adminlevel_governorate_code'])
                        #     try:
                        #         if report_type == 'live':
                        #             denominator = denominator_indicator.values_sections_partners_gov_live[
                        #                 key3] if key3 in denominator_indicator.values_sections_partners_gov_live else 0
                        #             numerator = numerator_indicator.values_sections_partners_gov_live[
                        #                 key3] if key3 in numerator_indicator.values_sections_partners_gov_live else 0
                        #         else:
                        #             denominator = denominator_indicator.values_sections_partners_gov[
                        #                 key3] if key3 in denominator_indicator.values_sections_partners_gov else 0
                        #             numerator = numerator_indicator.values_sections_partners_gov[
                        #                 key3] if key3 in numerator_indicator.values_sections_partners_gov else 0
                        #         denominator = denominator * denominator_multiplication
                        #         values_sections_partners_gov[key3] = numerator / denominator
                        #     except Exception:
                        #         values_sections_partners_gov[key3] = 0

                # for section in sections:
                #     if 'reporting_section' in section and section['reporting_section']:
                #         key4 = "{}-{}-{}".format(month, section['reporting_section'],partner['partner_id'])
                #         try:
                #             if report_type == 'live':
                #                 denominator = denominator_indicator.values_sections_partners_live[
                #                     key4] if key4 in denominator_indicator.values_sections_partners_live else 0
                #                 numerator = numerator_indicator.values_sections_partners_live[
                #                     key4] if key4 in numerator_indicator.values_sections_partners_live else 0
                #             else:
                #                 denominator = denominator_indicator.values_sections_partners[
                #                     key4] if key4 in denominator_indicator.values_sections_partners else 0
                #                 numerator = numerator_indicator.values_sections_partners[
                #                     key4] if key4 in numerator_indicator.values_sections_partners else 0
                #             denominator = denominator * denominator_multiplication
                #             values_sections_partners[key4] = numerator / denominator
                #         except Exception:
                #             values_sections_partners[key4] = 0

            if report_type == 'live':
                indicator.values_live[month] = values_month
                indicator.values_gov_live.update(values_gov)
                indicator.values_partners_live.update(values_partners)
                indicator.values_partners_gov_live.update(values_partners_gov)
                indicator.values_sections_live.update(values_sections)
                indicator.values_sections_partners_live.update(values_sections_partners)
                indicator.values_sections_gov_live.update(values_sections_gov)
                indicator.values_sections_partners_gov_live.update(values_sections_partners_gov)
            else:
                # if month == reporting_month:
                #     indicator.values_hpm[reporting_month] = values_month
                indicator.values[month] = values_month
                indicator.values_gov.update(values_gov)
                indicator.values_partners.update(values_partners)
                indicator.values_partners_gov.update(values_partners_gov)
                indicator.values_sections.update(values_sections)
                indicator.values_sections_partners.update(values_sections_partners)
                indicator.values_sections_gov.update(values_sections_gov)
                indicator.values_sections_partners_gov.update(values_sections_partners_gov)

        indicator.save()


#  todo to remove
def calculate_individual_indicators_values_11(ai_db):

    from internos.activityinfo.models import Indicator, ActivityReport

    # last_month = int(datetime.datetime.now().strftime("%m"))
    last_month = get_extraction_month(ai_db)

    reports = ActivityReport.objects.filter(database_id=ai_db.ai_id)
    if ai_db.is_funded_by_unicef:
        reports = reports.filter(funded_by='UNICEF')

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id).exclude(master_indicator=True)\
        .exclude(master_indicator_sub=True).only(
        'ai_indicator',
        'values_live',
        'values_gov_live',
        'values_partners_live')
    partners = reports.values('partner_id').distinct().order_by('partner_id')
    governorates = reports.values('location_adminlevel_governorate_code').distinct()
    governorates1 = reports.values('location_adminlevel_governorate_code').distinct()

    for indicator in indicators.iterator():
        qs_raw = ActivityReport.objects.raw(
            "SELECT id FROM activityinfo_activityreport "
            "WHERE indicator_id = %s AND funded_by = %s ",
            [indicator.ai_indicator, 'UNICEF'])
        try:
            count = qs_raw[0]
        except Exception as ex:
            # print(ex.message)
            continue

        for month in range(1, last_month):
            month = str(month)
            result = 0
            qs_raw = ActivityReport.objects.raw(
                "SELECT id, SUM(indicator_value) as indicator_value FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND indicator_id = %s AND funded_by = %s "
                "GROUP BY id",
                [month, indicator.ai_indicator, 'UNICEF'])
            try:
                result = qs_raw[0].indicator_value
            except Exception:
                continue

            # if month == reporting_month:
            #     indicator.values_hpm[reporting_month] = result
            indicator.values[str(month)] = result

            for gov1 in governorates1:
                value = 0
                key = "{}-{}".format(month, gov1['location_adminlevel_governorate_code'])

                qs_raw = ActivityReport.objects.raw(
                    "SELECT id, SUM(indicator_value) as indicator_value FROM activityinfo_activityreport "
                    "WHERE date_part('month', start_date) = %s AND indicator_id = %s AND funded_by = %s "
                    "AND location_adminlevel_governorate_code = %s "
                    "GROUP BY id",
                    [month, indicator.ai_indicator, 'UNICEF', gov1['location_adminlevel_governorate_code']])
                try:
                    value = qs_raw[0].indicator_value
                except Exception:
                    pass

                indicator.values_gov[str(key)] = value

            for partner in partners:
                value1 = 0
                key1 = "{}-{}".format(month, partner['partner_id'])

                qs_raw = ActivityReport.objects.raw(
                    "SELECT id, SUM(indicator_value) as indicator_value FROM activityinfo_activityreport "
                    "WHERE date_part('month', start_date) = %s AND indicator_id = %s AND funded_by = %s "
                    "AND partner_id = %s "
                    "GROUP BY id",
                    [month, indicator.ai_indicator, 'UNICEF', partner['partner_id']])
                try:
                    value1 = qs_raw[0].indicator_value
                except Exception:
                    continue

                indicator.values_partners[str(key1)] = value1

                for gov in governorates:
                    value2 = 0
                    key2 = "{}-{}-{}".format(month, partner['partner_id'], gov['location_adminlevel_governorate_code'])

                    qs_raw = ActivityReport.objects.raw(
                        "SELECT id, SUM(indicator_value) as indicator_value FROM activityinfo_activityreport "
                        "WHERE date_part('month', start_date) = %s AND indicator_id = %s AND funded_by = %s "
                        "AND partner_id = %s AND location_adminlevel_governorate_code = %s "
                        "GROUP BY id",
                        [month, indicator.ai_indicator, 'UNICEF', partner['partner_id'], gov['location_adminlevel_governorate_code']])
                    try:
                        value2 = qs_raw[0].indicator_value
                    except Exception:
                        pass

                    indicator.values_partners_gov[str(key2)] = value2

            indicator.save()


def calculate_individual_indicators_values_1(ai_db):
    from django.db import connection
    from internos.activityinfo.models import Indicator

    ai_id = str(ai_db.ai_id)

    if ai_db.is_current_extraction:
        last_month = get_current_extraction_month(ai_db) + 1
        # last_month = int(datetime.datetime.now().strftime("%m")) + 1
    else:
        last_month = get_extraction_month(ai_db)
        # last_month = int(datetime.datetime.now().strftime("%m"))

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id).exclude(master_indicator=True)\
        .exclude(master_indicator_sub=True).only(
        'ai_indicator',
        'values',
        'values_gov',
        'values_partners',
        'values_partners_gov',
        'values_sections',
        'values_sections_partners',
        'values_sections_gov',
        'values_sections_partners_gov',
    )

    rows_months = {}
    rows_partners = {}
    rows_govs = {}
    rows_partners_govs = {}
    rows_sections = {}
    rows_sections_partners = {}
    rows_sections_gov = {}
    rows_sections_partners_gov = {}

    values_hpm = {}
    cursor = connection.cursor()
    for month in range(1, last_month):
        month = str(month)

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id",
                [month, ai_id])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_months:
                rows_months[row[0]] = {}
            rows_months[row[0]][month] = row[1]

            # if month == reporting_month:
            #     if row[0] not in values_hpm:
            #         values_hpm[row[0]] = {}
            #     values_hpm[row[0]] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, location_adminlevel_governorate_code "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, location_adminlevel_governorate_code",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, location_adminlevel_governorate_code "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, location_adminlevel_governorate_code",
                [month, ai_id])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_govs:
                rows_govs[row[0]] = {}
            key = "{}-{}".format(month, row[2])
            rows_govs[row[0]][key] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, partner_id "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, partner_id",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, partner_id "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, partner_id",
                [month, ai_id])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_partners:
                rows_partners[row[0]] = {}
            key = "{}-{}".format(month, row[2])
            rows_partners[row[0]][key] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, location_adminlevel_governorate_code, partner_id "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, location_adminlevel_governorate_code, partner_id",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, location_adminlevel_governorate_code, partner_id "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, location_adminlevel_governorate_code, partner_id",
                [month, ai_id])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_partners_govs:
                rows_partners_govs[row[0]] = {}
            key = "{}-{}-{}".format(month, row[2], row[3])
            rows_partners_govs[row[0]][key] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, reporting_section",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, reporting_section",
                [month, ai_id])
        rows = cursor.fetchall()
        for row in rows:
            if row[0] not in rows_sections:
                rows_sections[row[0]] = {}
            key = "{}-{}".format(month, row[2])
            rows_sections[row[0]][key] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section ,partner_id "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, reporting_section ,partner_id",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section ,partner_id "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, reporting_section ,partner_id",
                [month, ai_id])
        rows = cursor.fetchall()
        for row in rows:
            if row[0] not in rows_sections_partners:
                rows_sections_partners[row[0]] = {}
            key = "{}-{}-{}".format(month, row[2], row[3])
            rows_sections_partners[row[0]][key] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section, "
                "location_adminlevel_governorate_code "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, reporting_section ,location_adminlevel_governorate_code",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section, "
                "location_adminlevel_governorate_code "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, reporting_section ,location_adminlevel_governorate_code",
                [month, ai_id])
        rows = cursor.fetchall()
        for row in rows:
            if row[0] not in rows_sections_gov:
                rows_sections_gov[row[0]] = {}
            key = "{}-{}-{}".format(month, row[2], row[3])
            rows_sections_gov[row[0]][key] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section, "
                "location_adminlevel_governorate_code, partner_id "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, reporting_section , partner_id ,location_adminlevel_governorate_code",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section , partner_id, location_adminlevel_governorate_code "
                "FROM activityinfo_activityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, reporting_section ,partner_id , location_adminlevel_governorate_code",
                [month, ai_id])
        rows = cursor.fetchall()
        for row in rows:
            if row[0] not in rows_sections_partners_gov:
                rows_sections_partners_gov[row[0]] = {}
            key = "{}-{}-{}-{}".format(month, row[2], row[3],row[4])
            rows_sections_partners_gov[row[0]][key] = row[1]


    for indicator in indicators.iterator():
        if indicator.ai_indicator in rows_months:
            indicator.values = rows_months[indicator.ai_indicator]

        # if indicator.ai_indicator in values_hpm:
        #     indicator.values_hpm[reporting_month] = values_hpm[indicator.ai_indicator]

        if indicator.ai_indicator in rows_partners:
            indicator.values_partners = rows_partners[indicator.ai_indicator]

        if indicator.ai_indicator in rows_govs:
            indicator.values_gov = rows_govs[indicator.ai_indicator]

        if indicator.ai_indicator in rows_partners_govs:
            indicator.values_partners_gov = rows_partners_govs[indicator.ai_indicator]

        if indicator.ai_indicator in rows_sections:
            indicator.values_sections = rows_sections[indicator.ai_indicator]

        if indicator.ai_indicator in rows_sections_gov:
            indicator.values_sections_gov = rows_sections_gov[indicator.ai_indicator]

        if indicator.ai_indicator in rows_sections_partners:
            indicator.values_sections_partners = rows_sections_partners[indicator.ai_indicator]

        if indicator.ai_indicator in rows_sections_partners_gov:
            indicator.values_sections_partners_gov = rows_sections_partners_gov[indicator.ai_indicator]

        indicator.save()


def calculate_individual_indicators_values_2(ai_db):
    from django.db import connection
    from internos.activityinfo.models import Indicator

    last_month = int(datetime.datetime.now().strftime("%m")) + 1
    # last_month = 13

    ai_id = str(ai_db.ai_id)
    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id).exclude(master_indicator=True)\
        .exclude(master_indicator_sub=True).only(
        'ai_indicator',
        'values_live',
        'values_gov_live',
        'values_partners_live',
        'values_partners_gov_live',
        'values_sections_live',
        'values_sections_partners_live',
        'values_sections_gov_live',
        'values_sections_partners_gov_live',
        'values_sections',
        'values_sections_partners',
        'values_sections_gov',
        'values_sections_partners_gov',
        'values_sections_live',
        'values_sections_partners_live',
        'values_sections_gov_live',
        'values_sections_partners_gov_live'
    )

    rows_months = {}
    rows_partners = {}
    rows_govs = {}
    rows_partners_govs = {}
    rows_sections = {}
    rows_sections_partners = {}
    rows_sections_gov = {}
    rows_sections_partners_gov = {}

    cursor = connection.cursor()
    for month in range(1, last_month):
        month = str(month)
        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id",
                [month, ai_id])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_months:
                rows_months[row[0]] = {}
            rows_months[row[0]][month] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, location_adminlevel_governorate_code "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, location_adminlevel_governorate_code",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, location_adminlevel_governorate_code "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, location_adminlevel_governorate_code",
                [month, ai_id])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_govs:
                rows_govs[row[0]] = {}
            key = "{}-{}".format(month, row[2])
            rows_govs[row[0]][key] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, partner_id "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, partner_id",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, partner_id "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, partner_id",
                [month, ai_id])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_partners:
                rows_partners[row[0]] = {}
            key = "{}-{}".format(month, row[2])
            rows_partners[row[0]][key] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, location_adminlevel_governorate_code, partner_id "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, location_adminlevel_governorate_code, partner_id",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, location_adminlevel_governorate_code, partner_id "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, location_adminlevel_governorate_code, partner_id",
                [month, ai_id])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_partners_govs:
                rows_partners_govs[row[0]] = {}
            key = "{}-{}-{}".format(month, row[2], row[3])
            rows_partners_govs[row[0]][key] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, reporting_section",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, reporting_section",
                [month, ai_id])
        rows = cursor.fetchall()
        for row in rows:
            if row[0] not in rows_sections:
                rows_sections[row[0]] = {}
            key = "{}-{}".format(month, row[2])
            rows_sections[row[0]][key] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section ,partner_id "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, reporting_section ,partner_id",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section ,partner_id "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, reporting_section ,partner_id",
                [month, ai_id])
        rows = cursor.fetchall()
        for row in rows:
            if row[0] not in rows_sections_partners:
                rows_sections_partners[row[0]] = {}
            key = "{}-{}-{}".format(month, row[2], row[3])
            rows_sections_partners[row[0]][key] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section ,location_adminlevel_governorate_code "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, reporting_section ,location_adminlevel_governorate_code",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section ,location_adminlevel_governorate_code "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, reporting_section ,location_adminlevel_governorate_code",
                [month, ai_id])
        rows = cursor.fetchall()
        for row in rows:
            if row[0] not in rows_sections_gov:
                rows_sections_gov[row[0]] = {}
            key = "{}-{}-{}".format(month, row[2], row[3])
            rows_sections_gov[row[0]][key] = row[1]

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section , "
                "location_adminlevel_governorate_code ,partner_id "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
                "GROUP BY indicator_id, reporting_section , partner_id ,location_adminlevel_governorate_code",
                [month, ai_id, 'UNICEF'])
        else:
            cursor.execute(
                "SELECT indicator_id, SUM(indicator_value) as indicator_value, reporting_section , partner_id ,location_adminlevel_governorate_code "
                "FROM activityinfo_liveactivityreport "
                "WHERE date_part('month', start_date) = %s AND database_id = %s "
                "GROUP BY indicator_id, reporting_section ,partner_id , location_adminlevel_governorate_code",
                [month, ai_id])
        rows = cursor.fetchall()
        for row in rows:
            if row[0] not in rows_sections_partners_gov:
                rows_sections_partners_gov[row[0]] = {}
            key = "{}-{}-{}-{}".format(month, row[2], row[3], row[4])
            rows_sections_partners_gov[row[0]][key] = row[1]

    for indicator in indicators.iterator():
        if indicator.ai_indicator in rows_months:
            indicator.values_live = rows_months[indicator.ai_indicator]

        if indicator.ai_indicator in rows_partners:
            indicator.values_partners_live = rows_partners[indicator.ai_indicator]

        if indicator.ai_indicator in rows_govs:
            indicator.values_gov_live = rows_govs[indicator.ai_indicator]

        if indicator.ai_indicator in rows_partners_govs:
            indicator.values_partners_gov_live = rows_partners_govs[indicator.ai_indicator]

        if indicator.ai_indicator in rows_sections:
            indicator.values_sections_live = rows_sections[indicator.ai_indicator]

        if indicator.ai_indicator in rows_sections_gov:
            indicator.values_sections_gov_live = rows_sections_gov[indicator.ai_indicator]

        if indicator.ai_indicator in rows_sections_partners:
            indicator.values_sections_partners_live = rows_sections_partners[indicator.ai_indicator]

        if indicator.ai_indicator in rows_sections_partners_gov:
            indicator.values_sections_partners_gov_live = rows_sections_partners_gov[indicator.ai_indicator]

        indicator.save()


def calculate_indicators_status(database):
    from internos.activityinfo.models import Indicator

    year_days = 365
    today = datetime.datetime.now()
    reporting_year = database.reporting_year
    beginning_year = datetime.datetime(int(reporting_year.year), 01, 01)
    # datetime.datetime(int(reporting_year.name), 01, 01)
    delta = today - beginning_year
    total_days = delta.days + 1
    days_passed_per = (total_days * 100) / year_days

    indicators = Indicator.objects.filter(activity__database__ai_id=database.ai_id)

    for indicator in indicators:
        cumulative_per = indicator.cumulative_per
        off_track = days_passed_per - 10
        over_target = days_passed_per + 10
        if cumulative_per < off_track:
            indicator.status = 'Off Track'
            indicator.status_color = 'badge-danger'
        elif cumulative_per > over_target:
            indicator.status = 'Over Track'
            indicator.status_color = 'badge-warning'
        else:
            indicator.status = 'On Track'
            indicator.status_color = 'badge-success'

        cumulative_per = indicator.cumulative_per_sector
        off_track = days_passed_per - 10
        over_target = days_passed_per + 10
        if cumulative_per < off_track:
            indicator.status_sector = 'Off Track'
            indicator.status_color_sector = 'badge-danger'
        elif cumulative_per > over_target:
            indicator.status_sector = 'Over Track'
            indicator.status_color_sector = 'badge-warning'
        else:
            indicator.status_sector = 'On Track'
            indicator.status_color_sector = 'badge-success'

        indicator.save()

    return indicators.count()


def assign_main_master_indicator():
    from internos.activityinfo.models import Indicator

    # Level 1
    top_indicators = Indicator.objects.filter(master_indicator=True).only(
        'sub_indicators',
    )

    for indicator in top_indicators.iterator():
        sub_indicators = indicator.sub_indicators.all()
        sub_indicators.update(main_master_indicator=indicator)

    # Level 2
    top_indicators1 = Indicator.objects.filter(master_indicator_sub=True).only(
        'sub_indicators',
        'main_master_indicator',
    )

    for indicator in top_indicators1.iterator():
        sub_indicators = indicator.sub_indicators.all()
        sub_indicators.update(main_master_indicator=indicator.main_master_indicator)

    # Level 3
    top_indicators2 = Indicator.objects.filter(master_indicator_sub_sub=True).only(
        'sub_indicators',
        'main_master_indicator',
    )

    for indicator in top_indicators2.iterator():
        sub_indicators = indicator.sub_indicators.all()
        sub_indicators.update(main_master_indicator=indicator.main_master_indicator)


def assign_support_disability_master_indicator():
    from internos.activityinfo.models import Indicator

    # Level 4
    sub_indicators = Indicator.objects.filter(tag_disability__isnull=False)
    sub_indicators.update(support_disability=True)

    # Level 3
    sub_indicators2 = Indicator.objects.filter(master_indicator_sub_sub=True, sub_indicators__support_disability=True)
    sub_indicators2.update(support_disability=True)

    # Level 2
    sub_indicators3 = Indicator.objects.filter(master_indicator_sub=True, sub_indicators__support_disability=True)
    sub_indicators3.update(support_disability=True)

    # Level 1
    top_indicators = Indicator.objects.filter(master_indicator=True)

    for indicator in top_indicators.iterator():
        sub_indicators = indicator.sub_indicators.filter(tag_disability__isnull=False)
        sub_indicators.update(support_disability=True)

    # Level 2
    top_indicators1 = Indicator.objects.filter(master_indicator_sub=True).only(
        'sub_indicators',
        'support_disability',
    )

    for indicator in top_indicators1.iterator():
        sub_indicators = indicator.sub_indicators.all()
        sub_indicators.update(support_disability=indicator.support_disability)

    # Level 3
    top_indicators2 = Indicator.objects.filter(master_indicator_sub_sub=True).only(
        'sub_indicators',
        'support_disability',
    )

    for indicator in top_indicators2.iterator():
        sub_indicators = indicator.sub_indicators.all()
        sub_indicators.update(support_disability=indicator.support_disability)


def load_reporting_map(partner=None, governorate=None, caza=None, donor=None, indicator=None, ai_id=None):
    from django.db import connection

    params = [str(ai_id), ]

    cursor = connection.cursor()
    queryset = "SELECT DISTINCT ar.site_id, ar.location_name, ar.location_longitude, ar.location_latitude, " \
               "ar.indicator_units, ar.location_adminlevel_governorate, ar.location_adminlevel_caza, " \
               "ar.location_adminlevel_caza_code, ar.location_adminlevel_cadastral_area, " \
               "ar.location_adminlevel_cadastral_area_code, ar.partner_label, ai.name AS indicator_name, " \
               "ai.cumulative_values ->> 'months'::text AS cumulative_value " \
               "FROM public.activityinfo_indicator ai " \
               "INNER JOIN public.activityinfo_activityreport ar ON ai.id = ar.ai_indicator_id " \
               "WHERE ar.database_id = %s "

    if partner:
        params.append(str(partner))
        queryset += "AND partner_id = %s "
    if governorate:
        params.append(str(governorate))
        queryset += "AND location_adminlevel_governorate_code = %s "
    if caza:
        params.append(str(caza))
        queryset += "AND location_adminlevel_caza_code = %s "
    if indicator:
        params.append(str(indicator))
        params.append(str(indicator))
        queryset += "AND (ai.main_master_indicator_id = %s OR ai.id = %s) "

    cursor.execute(queryset, params)
    return cursor.fetchall()


def update_partner_data(ai_db):
    from .client import ActivityInfoClient
    from .models import Partner

    client = ActivityInfoClient(ai_db.username, ai_db.password)

    dbs = client.get_databases()
    db_ids = [db['id'] for db in dbs]
    if ai_db.ai_id not in db_ids:
        raise Exception(
            'DB with ID {} not found in ActivityInfo'.format(
                ai_db.ai_id
            ))

    db_info = client.get_database(ai_db.ai_id)

    objects = 0
    try:
        if 'databasePartners' in db_info:
            partners = db_info['databasePartners']
        else:
            partners = db_info['partners']

        for partner in partners:
            try:
                ai_partner = Partner.objects.get(ai_id=partner['id'])
            except Partner.DoesNotExist:
                ai_partner = Partner(ai_id=partner['id'])
                objects += 1
            ai_partner.name = partner['name']
            ai_partner.full_name = partner['fullName']
            ai_partner.database = ai_db
            ai_partner.save()

    except Exception as e:
        raise e

    return objects


def update_indicator_data(ai_db, ai_field_name, field_name):
    from internos.activityinfo.client import ActivityInfoClient
    from internos.activityinfo.models import Indicator

    client = ActivityInfoClient(ai_db.username, ai_db.password)

    dbs = client.get_databases()
    db_ids = [db['id'] for db in dbs]
    if ai_db.ai_id not in db_ids:
        raise Exception(
            'DB with ID {} not found in ActivityInfo'.format(
                ai_db.ai_id
            ))

    db_info = client.get_database(ai_db.ai_id)

    objects = 0
    try:

        for activity in db_info['activities']:
            for indicator in activity['indicators']:
                try:
                    ai_indicator = Indicator.objects.get(ai_id=indicator['id'])
                except Indicator.DoesNotExist:
                    continue

                objects += 1
                setattr(ai_indicator, ai_field_name, indicator[field_name])
                ai_indicator.save()

    except Exception as e:
        raise e

    return objects


def update_hpm_report():
    update_indicators_hpm_data()
    # calculate_indicators_cumulative_hpm()
    # calculate_indicators_tags_hpm()


def update_indicators_hpm_data():
    from internos.activityinfo.models import Indicator

    current_month = datetime.datetime.now().month
    current_year = datetime.datetime.now().year

    indicators = Indicator.objects.filter(activity__database__reporting_year__year=current_year ,hpm_indicator=True).only(
        'values',
        'values_hpm',
        'cumulative_values',
        'cumulative_values_hpm',
    )
    months =[]
    for m in range(1,current_month):
        months.append(m)

    for indicator in indicators.iterator():
        value = 0
        for m in months:
            if str(m) in indicator.values:
                # if str(m) in indicator.values_hpm:
                #    continue
                # else:
                value = indicator.values[str(m)]
                indicator.values_hpm[str(m)] = value
        indicator.save()
        # code was added to calculate periodic data each 3 months
        # cum_value = 0
        # cum_value2 = 0
        # cum_value3 = 0
        # cum_value4 = 0
        # for m in months:
        #     if str(m) in indicator.values:
        #         if 0 < m < 4:
        #             cum_value += indicator.values[str(m)]
        #         if 3 < m < 7:
        #             cum_value2 += indicator.values[str(m)]
        #         if 6 < m < 10:
        #             cum_value3 += indicator.values[str(m)]
        #         if 9 < m < 13:
        #             cum_value4 += indicator.values[str(m)]
        #
        # indicator.cumulative_values_hpm['3'] = cum_value
        # indicator.cumulative_values_hpm['6'] = cum_value2
        # indicator.cumulative_values_hpm['9'] = cum_value3
        # indicator.cumulative_values_hpm['12'] = cum_value4
        # indicator.save()


def update_hpm_table_docx(indicators, month, month_name, filename,reporting_year,type=None):
    
    from docx import Document
    from internos.activityinfo.templatetags.util_tags import get_hpm_indicator_data_new , get_hpm_sub_indicators
    from internos.activityinfo.models import Database , Indicator
    from docx.shared import Pt

    path = os.path.dirname(os.path.abspath(__file__))
    path2file = path+'/AIReports/HPM Table Template ' + str(reporting_year) +'.docx'

    document = Document(path2file)
    # month = month - 1
    # document.paragraphs[0].runs[1].text = month_name
    replacements = {
        'month_name': month_name ,
    }
    #
    # for paragraph in document.paragraphs:
    #     for key in replacements:
    #         paragraph.text = paragraph.text.replace(key, replacements[key])
    # document.paragraphs[0].runs[0].style
    #
    # for paragraph in document.tables[0].rows[0].cells[1].paragraphs:
    #     for key in replacements:
    #         paragraph.text = paragraph.text.replace(key, replacements[key])
    if month == 1:
        document.paragraphs[0].add_run('{} {} {} {}'.format('HPM Table | Data of', month_name ,'|',reporting_year))

        document.tables[0].rows[0].cells[1].paragraphs[0].runs[0].text='{} {} {} {} {}'.format('SUMMARY OF PROGRAMME RESULTS |', month_name,'|',reporting_year,'SITREP - LEBANON')
    else:
        document.paragraphs[0].add_run('{} {} {} {}'.format('HPM Table | Data of January to ', month_name, '|', reporting_year))

        document.tables[0].rows[0].cells[1].paragraphs[0].runs[0].text = '{} {} {} {} {}'.format('SUMMARY OF PROGRAMME RESULTS | January to', month_name, '|', reporting_year, 'SITREP - LEBANON')

    databases = Database.objects.filter(reporting_year__name=reporting_year).exclude(hpm_sequence=0).order_by('hpm_sequence')
    # Note that sequence is important here since it affects the code in getting databases.
    row_num = 2

    edu_list=[]
    edu_indicators = Indicator.objects.filter(activity__database=databases[0],hpm_indicator=True,master_indicator=True).order_by('sequence')
    try:
        sub_indicators = get_hpm_sub_indicators(edu_indicators[0].id)
        for item in sub_indicators:
            edu_list.append(item['id'])
        edu_list.append(edu_indicators[0].id)
        sub_indicators = get_hpm_sub_indicators(edu_indicators[1].id)
        for item in sub_indicators:
            edu_list.append(item['id'])
        edu_list.append(edu_indicators[1].id)
    except Exception as ex:
        logger.error(ex.message)

    education_ids = edu_list
    # education_ids = [7019,7020,7021,6959,6958,6955,6960]
    # Education 1
    key = 0 
    for id in education_ids:
        
        indicator1 = get_hpm_indicator_data_new(id, month,type)
        document.tables[0].rows[row_num].cells[1].paragraphs[0].runs[0].text = str(indicator1['hpm_label'])
        document.tables[0].rows[row_num].cells[3].paragraphs[0].runs[0].text = str(indicator1['target_sector'])
        document.tables[0].rows[row_num].cells[6].paragraphs[0].runs[0].text = str(indicator1['target'])
        document.tables[0].rows[row_num].cells[7].paragraphs[0].runs[0].text = str(indicator1['cumulative'])
        if indicator1['has_hpm_note']:
            key = key + 1
            super_text = document.tables[0].rows[row_num].cells[7].paragraphs[0].add_run(str(key))
            super_text.font.superscript = True
            super_text.font.name = 'Calibri'
            super_text.font.size = Pt(9)
    
        
        document.tables[0].rows[row_num].cells[8].paragraphs[0].runs[0].text = str(indicator1['report_change'])
        row_num = row_num + 1

    ## CP
    Child_indicators = Indicator.objects.filter(activity__database=databases[1], hpm_indicator=True,master_indicator=True).order_by('sequence')

    row_num = row_num + 1
    # #child_protection_ids = [6972, 6990, 6946]
    key = 0 
    for indicator in Child_indicators:
        Child_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        key = key + 1
        document.tables[0].rows[row_num].cells[1].paragraphs[0].runs[0].text = str(Child_1['hpm_label'])
        document.tables[0].rows[row_num].cells[3].paragraphs[0].runs[0].text = str(Child_1['target_sector'])
        document.tables[0].rows[row_num].cells[6].paragraphs[0].runs[0].text = str(Child_1['target'])
        document.tables[0].rows[row_num].cells[7].paragraphs[0].runs[0].text = str(Child_1['cumulative'])
        if Child_1['has_hpm_note']:
            
            super_text = document.tables[0].rows[row_num].cells[7].paragraphs[0].add_run(str(key))
            super_text.font.superscript = True
            super_text.font.name = 'Calibri'
            super_text.font.size = Pt(9)
            
        document.tables[0].rows[row_num].cells[8].paragraphs[0].runs[0].text = str(Child_1['report_change'])
        row_num = row_num + 1

    SVBG_indicators = Indicator.objects.filter(activity__database=databases[8], hpm_indicator=True, master_indicator=True).order_by('sequence')
    
    for indicator in SVBG_indicators:
        Child_3 = get_hpm_indicator_data_new(indicator.id, month,type)
        key = key + 1
        document.tables[0].rows[row_num].cells[1].paragraphs[0].runs[0].text = str(Child_3['hpm_label'])
        document.tables[0].rows[row_num].cells[3].paragraphs[0].runs[0].text = str(Child_3['target_sector'])
        document.tables[0].rows[row_num].cells[6].paragraphs[0].runs[0].text = str(Child_3['target'])
        document.tables[0].rows[row_num].cells[7].paragraphs[0].runs[0].text = str(Child_3['cumulative'])
        if Child_3['has_hpm_note']:
            
            super_text = document.tables[0].rows[row_num].cells[7].paragraphs[0].add_run(str(key))
            super_text.font.superscript = True
            super_text.font.name = 'Calibri'
            super_text.font.size = Pt(9)
            
        document.tables[0].rows[row_num].cells[8].paragraphs[0].runs[0].text = str(Child_3['report_change'])
        row_num = row_num + 1


   #  #
   #  # # WASH
    Wash_indicators = Indicator.objects.filter(activity__database=databases[2], hpm_indicator=True,
                                                master_indicator=True).order_by('sequence')
    row_num = row_num + 1
    key = 0 
    # wash_ids = [6995, 6996, 6999,6997]
    for indicator in Wash_indicators:
        wash_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        key = key + 1
        document.tables[0].rows[row_num].cells[1].paragraphs[0].runs[0].text = str(wash_1['hpm_label'])
        document.tables[0].rows[row_num].cells[3].paragraphs[0].runs[0].text = str(wash_1['target_sector'])
        document.tables[0].rows[row_num].cells[6].paragraphs[0].runs[0].text = str(wash_1['target'])
        document.tables[0].rows[row_num].cells[7].paragraphs[0].runs[0].text = str(wash_1['cumulative'])
        if wash_1['has_hpm_note']:
            
            super_text = document.tables[0].rows[row_num].cells[7].paragraphs[0].add_run(str(key))
            super_text.font.superscript = True
            super_text.font.name = 'Calibri'
            super_text.font.size = Pt(9)
        document.tables[0].rows[row_num].cells[8].paragraphs[0].runs[0].text = str(wash_1['report_change'])
        row_num = row_num + 1

   #  # # H&N
    health_indicators = Indicator.objects.filter(activity__database=databases[3], hpm_indicator=True,
                                               master_indicator=True).order_by('sequence')
    row_num = row_num + 1
    key = 0 
    # health_ids = [6941,6942,6940]
    for indicator in health_indicators:
        health_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        key = key + 1
        document.tables[0].rows[row_num].cells[1].paragraphs[0].runs[0].text = str(health_1['hpm_label'])
        document.tables[0].rows[row_num].cells[3].paragraphs[0].runs[0].text = str(health_1['target_sector'])
        document.tables[0].rows[row_num].cells[6].paragraphs[0].runs[0].text = str(health_1['target'])
        document.tables[0].rows[row_num].cells[7].paragraphs[0].runs[0].text = str(health_1['cumulative'])
        if health_1['has_hpm_note']:
            
            super_text = document.tables[0].rows[row_num].cells[7].paragraphs[0].add_run(str(key))
            super_text.font.superscript = True
            super_text.font.name = 'Calibri'
            super_text.font.size = Pt(9)
        document.tables[0].rows[row_num].cells[8].paragraphs[0].runs[0].text = str(health_1['report_change'])
        row_num = row_num + 1

   #  # # Y&A
   #  youth_ids = [6902,6904,6922,6928]
    youth_indicators = Indicator.objects.filter(activity__database=databases[4], hpm_indicator=True,
                                                master_indicator=True).order_by('sequence')
    row_num = row_num + 1
    key = 0 
    for indicator in youth_indicators:
        youth_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        key = key + 1
        document.tables[0].rows[row_num].cells[1].paragraphs[0].runs[0].text = str(youth_1['hpm_label'])
        document.tables[0].rows[row_num].cells[3].paragraphs[0].runs[0].text = str(youth_1['target_sector'])
        document.tables[0].rows[row_num].cells[6].paragraphs[0].runs[0].text = str(youth_1['target'])
        document.tables[0].rows[row_num].cells[7].paragraphs[0].runs[0].text = str(youth_1['cumulative'])
        if youth_1['has_hpm_note']:
            
            super_text = document.tables[0].rows[row_num].cells[7].paragraphs[0].add_run(str(key))
            super_text.font.superscript = True
            super_text.font.name = 'Calibri'
            super_text.font.size = Pt(9)
        document.tables[0].rows[row_num].cells[8].paragraphs[0].runs[0].text = str(youth_1['report_change'])
        row_num = row_num + 1

   #  # # SP
    # sp_ids = [6910,6911]
    sp_indicators = Indicator.objects.filter(activity__database=databases[5], hpm_indicator=True,
                                                master_indicator=True).order_by('sequence')
    row_num = row_num + 1
    key = 0 
    for indicator in sp_indicators:
        sp_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        document.tables[0].rows[row_num].cells[1].paragraphs[0].runs[0].text = str(sp_1['hpm_label'])
        document.tables[0].rows[row_num].cells[3].paragraphs[0].runs[0].text = str(sp_1['target_sector'])
        document.tables[0].rows[row_num].cells[6].paragraphs[0].runs[0].text = str(sp_1['target'])
        if sp_1['is_cumulative']:
            document.tables[0].rows[row_num].cells[7].paragraphs[0].runs[0].text = str(sp_1['cumulative'])
            document.tables[0].rows[row_num].cells[8].paragraphs[0].runs[0].text = str(sp_1['report_change'])
        else:
            document.tables[0].rows[row_num].cells[7].paragraphs[0].runs[0].text = str(sp_1['highest'])
            document.tables[0].rows[row_num].cells[8].paragraphs[0].runs[0].text = str(sp_1['highest_change'])
        row_num = row_num + 1
   #  # # C4D

    # C4D_ids = [6917]
    C4D_indicators = Indicator.objects.filter(activity__database=databases[6], hpm_indicator=True,
                                             master_indicator=True).order_by('sequence')
    row_num = row_num + 1
    key = 0 
    for indicator in C4D_indicators:
        C4D_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        key = key + 1
        document.tables[0].rows[row_num].cells[1].paragraphs[0].runs[0].text = str(C4D_1['hpm_label'])
        document.tables[0].rows[row_num].cells[3].paragraphs[0].runs[0].text = str(C4D_1['target_sector'])
        document.tables[0].rows[row_num].cells[6].paragraphs[0].runs[0].text = str(C4D_1['target'])
        document.tables[0].rows[row_num].cells[7].paragraphs[0].runs[0].text = str(C4D_1['cumulative'])
        if C4D_1['has_hpm_note']:
            
            super_text = document.tables[0].rows[row_num].cells[7].paragraphs[0].add_run(str(key))
            super_text.font.superscript = True
            super_text.font.name = 'Calibri'
            super_text.font.size = Pt(9)
            
        document.tables[0].rows[row_num].cells[8].paragraphs[0].runs[0].text = str(C4D_1['report_change'])
        row_num = row_num + 1

   #  # # PPL
   #
    # PPL_ids = [6819,6820,6822,6823,6832,6833]
    PPL_indicators = Indicator.objects.filter(activity__database=databases[7], hpm_indicator=True,
                                              master_indicator=True).order_by('sequence')
    row_num = row_num + 1
    key = 0 

    for indicator in PPL_indicators:
        key = key + 1
        PPL_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        document.tables[0].rows[row_num].cells[1].paragraphs[0].runs[0].text = str(PPL_1['hpm_label'])
        document.tables[0].rows[row_num].cells[3].paragraphs[0].runs[0].text = str(PPL_1['target_sector'])
        document.tables[0].rows[row_num].cells[6].paragraphs[0].runs[0].text = str(PPL_1['target'])
        document.tables[0].rows[row_num].cells[7].paragraphs[0].runs[0].text = str(PPL_1['cumulative'])
        if PPL_1['has_hpm_note']:
            
            super_text = document.tables[0].rows[row_num].cells[7].paragraphs[0].add_run(str(key))
            super_text.font.superscript = True
            super_text.font.name = 'Calibri'
            super_text.font.size = Pt(9)
            
        document.tables[0].rows[row_num].cells[8].paragraphs[0].runs[0].text = str(PPL_1['report_change'])
        
        row_num = row_num + 1
   #  # # Footnotes

    row_num = row_num + 1
    # last_row = document.tables[0].rows[row_num]
    # table = document.add_table(1, 2)
    # table.add_row()

   ### Education
    index = 0 
    for id in education_ids:
        
        indicator_1 = get_hpm_indicator_data_new(id, month,type)
        if indicator_1['has_hpm_note']:
            row_num = row_num + 1
            index = index + 1
            
            document.tables[0].rows[row_num].cells[0].paragraphs[0].runs[0].text = databases[0].hpm_label+" ("+str(index)+")"
            
            if indicator_1['cumulative'] != "0":
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = ' {}: {}%, {}: {}%. '.format(
                    'Boys', indicator_1['male'],'Girls', indicator_1['female']), indicator_1['comment']
            else:
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = "No figures are reported yet"

    ### cp
    index = 0 
    for indicator in Child_indicators:
        index = index + 1
        Child_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        if Child_1['has_hpm_note']:
             row_num = row_num + 1
             
             document.tables[0].rows[row_num].cells[0].paragraphs[0].runs[0].text = databases[1].hpm_label+" ("+str(index)+")"
             if Child_1['cumulative'] != "0":
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = '{}: {}%, {}: {}%. '.format(
                    'Boys', Child_1['male'],'Girls', Child_1['female']), Child_1['comment']
             else:
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = "No figures are reported yet"

    # SVBG merged with CP
    
    SVBG_indicators = Indicator.objects.filter(activity__database=databases[8], hpm_indicator=True,
                                               master_indicator=True).order_by('sequence')
    for indicator in SVBG_indicators:
        index = index + 1
        Child_3 = get_hpm_indicator_data_new(indicator.id, month,type)
        if Child_3['has_hpm_note']:
            row_num = row_num + 1
            
            document.tables[0].rows[row_num].cells[0].paragraphs[0].runs[0].text = '{} ({})'.format(databases[1].hpm_label,str(index))
            if Child_3['cumulative'] != "0":
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = ' {}: {}%, {}: {}%. '.format(
                    'Boys', Child_3['male'],'Girls', Child_3['female']), Child_3['comment']
            else:
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = "No figures are reported yet"

    #  # #  WASH
    index = 0 
    for indicator in Wash_indicators:
        index = index + 1
        wash_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        if wash_1['has_hpm_note']:
            row_num = row_num + 1
            
            document.tables[0].rows[row_num].cells[0].paragraphs[0].runs[0].text = databases[2].hpm_label+" ("+str(index)+")"
            if wash_1['cumulative'] != "0":
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[
                    0].text = '{}: {}%, {}: {}%, {}: {}%, {}: {}%. '.format('Boys', wash_1['boys'], 'Girls',
                                                                                wash_1['girls'],'Female',
                                                                                wash_1['female'],'Male',
                                                                                wash_1['male']), wash_1['comment']
            else:
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = "No figures are reported yet"

    #  # Health
    index = 0 
    for indicator in health_indicators:
        index = index + 1
        health_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        if health_1['has_hpm_note']:
            row_num = row_num + 1
            
            document.tables[0].rows[row_num].cells[0].paragraphs[0].runs[0].text = databases[3].hpm_label+" ("+str(index)+")"
            if health_1['cumulative'] != "0":
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = '{}: {}%, {}: {}%. '.format(
                    'Male', health_1['male'],'Female', health_1['female']), health_1['comment']
            else:
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = "No figures are reported yet"

    #  # #  Y&A
    index = 0 
    for indicator in youth_indicators:
        index = index + 1
        youth_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        if youth_1['has_hpm_note']:
            row_num = row_num + 1
            
            document.tables[0].rows[row_num].cells[0].paragraphs[0].runs[0].text = databases[4].hpm_label+" ("+str(index)+")"
            if youth_1['cumulative'] != "0":
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = '{}: {}%, {}: {}%. '.format(
                    'Male', youth_1['male'],'Female', youth_1['female']), youth_1['comment']
            else:
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = "No figures are reported yet"

    ## Social Policy
    index = 0 
    for indicator in sp_indicators:
        index = index + 1
        sp_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        if sp_1['has_hpm_note']:
            row_num = row_num + 1
            
            document.tables[0].rows[row_num].cells[0].paragraphs[0].runs[0].text = databases[5].hpm_label+" ("+str(index)+")"
            if sp_1['cumulative'] != "0":
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = '{}: {}%, {}: {}%. '.format(
                    'Boys', sp_1['male'], 'Girls', sp_1['female']), sp_1['comment']
            else:
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = "No figures are reported yet"

    ## C4D
    index = 0 
    for indicator in C4D_indicators:
        index = index + 1    
        C4D_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        if C4D_1['has_hpm_note']:
            row_num = row_num + 1
            
            document.tables[0].rows[row_num].cells[0].paragraphs[0].runs[0].text = databases[6].hpm_label+" ("+str(index)+")"
            if C4D_1['cumulative'] != "0":
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = '{}: {}%, {}: {}%. '.format(
                    'Boys', C4D_1['male'], 'Girls', C4D_1['female']), C4D_1['comment']
            else:
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = "No figures are reported yet"

    # #  PPL
    index = 0 
    for indicator in PPL_indicators:
        index = index + 1
        PPL_1 = get_hpm_indicator_data_new(indicator.id, month,type)
        if PPL_1['has_hpm_note']:
            row_num = row_num + 1
            
            document.tables[0].rows[row_num].cells[0].paragraphs[0].runs[0].text = databases[7].hpm_label+" ("+str(index)+")"
            if PPL_1['cumulative'] != "0":
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = '{}: {}%, {}: {}%. '.format(
                    'Boys', PPL_1['male'], 'Girls', PPL_1['female']), PPL_1['comment']
            else:
                document.tables[0].rows[row_num].cells[2].paragraphs[0].runs[0].text = "No figures are reported yet"

    path2file2 = '{}/{}/{}'.format(path, 'AIReports', filename)
    document.save(path2file2)
    return path2file2


def update_hpm_table_docx1(indicators, month, filename):

    # import sys
    # import docx
    from docx import Document
    from internos.activityinfo.templatetags.util_tags import get_indicator_cumulative, get_indicator_hpm_data

    path = os.path.dirname(os.path.abspath(__file__))
    path2file = path+'/AIReports/HPM Table template 2019.docx'
    # path2file = path+'/AIReports/test.docx'

    document = Document(path2file)

    table = document.tables[0]

    for indicator in indicators:
        # cum_key = '#{}_cum#'.format(indicator.id)
        cum_key = str(indicator.id)
        # diff_key = '#{}_diff#'.format(indicator.id)
        diff_key = str(indicator.id)

        for row in range(0, 40):
            # try:
            p = table.rows[row].cells[6].paragraphs[0]
            # for p in table.rows[row].cells[6].paragraphs:
            if cum_key in p.text:
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if cum_key in inline[i].text:
                        document.tables[0].rows[row].cells[6].paragraphs[0].runs[i].text = get_indicator_cumulative(indicator, month)

            p = table.rows[row].cells[7].paragraphs[0]
            # for p in table.rows[row].cells[7].paragraphs:
            # print(diff_key)
            # if diff_key in p.text:
            #     print(p.text)

            inline = p.runs
            # print(inline)
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                # if diff_key == '1504':
                # print(inline[i].text)
                if diff_key in inline[i].text:
                    document.tables[0].rows[row].cells[7].paragraphs[0].runs[i].text = 0

            # except Exception as ex:
                # print(ex.message)
                # continue

    path2file2 = '{}/{}/{}'.format(path, 'AIReports', filename)
    document.save(path2file2)
    return path2file2


def calculate_internal_indicators_values(ai_db,indicator_id):
    from django.db import connection
    from internos.activityinfo.models import Indicator

    # last_month = int(datetime.datetime.now().strftime("%m"))
    # last_month = 13
    last_month = get_extraction_month(ai_db)

    indicators = Indicator.objects.filter(id=indicator_id).only(
        'ai_indicator',
        'values',
        'values_gov',
        'values_partners',
        'values_partners_gov'
    )
    rows_months = {}
    rows_partners = {}
    rows_govs = {}
    rows_partners_govs = {}

    cursor = connection.cursor()
    for month in range(1, last_month):
        month = str(month)
        cursor.execute(
            "SELECT indicator_id, SUM(indicator_value) as indicator_value "
            "FROM activityinfo_activityreport "
            "WHERE  date_part('month', start_date) = %s AND database_id = %s "
            "GROUP BY indicator_id",
            [month, ai_db])

        rows = cursor.fetchall()
        if rows:
            for row in rows:
                if row[0] not in rows_months:
                    rows_months[row[0]] = {}
                rows_months[row[0]][month] = row[1]

        cursor.execute(
            "SELECT indicator_id, SUM(indicator_value) as indicator_value, location_adminlevel_governorate_code "
            "FROM activityinfo_activityreport "
            "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
            "GROUP BY indicator_id, location_adminlevel_governorate_code",
            [month, ai_db, 'UNICEF'])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_govs:
                rows_govs[row[0]] = {}
            key = "{}-{}".format(month, row[2])
            rows_govs[row[0]][key] = row[1]

        cursor.execute(
            "SELECT indicator_id, SUM(indicator_value) as indicator_value, partner_id "
            "FROM activityinfo_activityreport "
            "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
            "GROUP BY indicator_id, partner_id",
            [month, ai_db, 'UNICEF'])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_partners:
                rows_partners[row[0]] = {}
            key = "{}-{}".format(month, row[2])
            rows_partners[row[0]][key] = row[1]


        cursor.execute(
            "SELECT indicator_id, SUM(indicator_value) as indicator_value, location_adminlevel_governorate_code, partner_id "
            "FROM activityinfo_activityreport "
            "WHERE date_part('month', start_date) = %s AND database_id = %s AND funded_by = %s "
            "GROUP BY indicator_id, location_adminlevel_governorate_code, partner_id",
            [month, ai_db, 'UNICEF'])

        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_partners_govs:
                rows_partners_govs[row[0]] = {}
            key = "{}-{}-{}".format(month, row[2], row[3])
            rows_partners_govs[row[0]][key] = row[1]

    for indicator in indicators.iterator():
        if str(indicator.id) in rows_months:
            indicator.values = rows_months[str(indicator.id)]
            indicator.values_gov = rows_govs[str(indicator.id)]
            indicator.values_partners = rows_partners[str(indicator.id)]
            indicator.values_partners_gov = rows_partners_govs[str(indicator.id)]

        indicator.save()


def calculate_internal_cumulative_results(ai_db,indicator_id):
    from django.db import connection
    from internos.activityinfo.models import Indicator

    indicators = Indicator.objects.filter(id=indicator_id).only(
        'id',
        'values',
        'values_gov',
        'values_partners',
        'values_partners_gov',
        'cumulative_values',
    )

    rows_data = {}
    cursor = connection.cursor()
    cursor.execute(
        "SELECT distinct ai.id, ai.ai_indicator, aa.id, aa.name, ai.values, ai.values_gov , ai.values_partners, "
        "ai.values_partners_gov ,ai.none_ai_indicator"
        " FROM public.activityinfo_indicator ai, public.activityinfo_activity aa "
        "WHERE ai.activity_id = aa.id AND ai.none_ai_indicator='True' AND aa.database_id = %s",[ai_db])

    rows = cursor.fetchall()

    for row in rows:
        rows_data[row[0]] = row

    for indicator in indicators.iterator():

        values_month = {}
        values_partners = {}
        values_gov = {}
        values_partners_gov = {}

        if indicator.id in rows_data:

            indicator_values = rows_data[indicator.id]
            values = indicator_values[4]  # values
            values1 = indicator_values[5]  # values_gov
            values2 = indicator_values[6]  # values_partners
            values3 = indicator_values[7]  # values_partners_gov

            c_value = 0
            for key, value in values.items():
                c_value += value
                values_month = c_value

            for key, value in values1.items():
                keys = key.split('-')
                gov = keys[1]
                if gov in values_gov:
                    values_gov[gov] = values_gov[gov] + value
                else:
                    values_gov[gov] = value

            for key, value in values2.items():
                keys = key.split('-')
                partner = keys[1]
                if partner in values_partners:
                    values_partners[partner] = values_partners[partner] + value
                else:
                    values_partners[partner] = value

            for key, value in values3.items():
                keys = key.split('-')
                gov_partner = '{}-{}'.format(keys[1], keys[2])
                if gov_partner in values_partners_gov:
                    values_partners_gov[gov_partner] = values_partners_gov[gov_partner] + value
                else:
                    values_partners_gov[gov_partner] = value

            indicator.cumulative_values = {
                'months': values_month,
                'govs': values_gov,
                'partners':values_partners,
                'partners_govs': values_partners_gov

            }
            indicator.save()

    return indicators.count()


def get_partners_list(ai_db, sections=None, govs=None, months=None, report_type=None):
    from django.db import connection

    result = {}
    cursor = connection.cursor()
    where_condition = ""
    funded_condition = ""

    if ai_db.is_funded_by_unicef:
        funded_condition = " AND funded_by = 'UNICEF' "
    if govs:
        govs_list = ", ".join("'" + str(n) + "'" for n in govs)
        where_condition = " and location_adminlevel_governorate_code in (" + govs_list + ")"

    if sections:
        section_list = ", ".join("'" + str(n) + "'" for n in sections)
        where_condition += " AND reporting_section in (" + section_list + ")"

    if months:
        month_list = ", ".join("'" + str(n) + "'" for n in months)
        where_condition += " AND date_part('month', start_date) in (" + month_list + ")"

    query_condition = "{}'{}'".format(" WHERE database_id =", str(ai_db.ai_id)) + funded_condition + where_condition

    if report_type == 'live':
        cursor.execute(
                "SELECT DISTINCT partner_id, partner_label "
                "FROM activityinfo_liveactivityreport "
                 + query_condition)
    else:
        cursor.execute(
                "SELECT DISTINCT partner_id, partner_label "
                "FROM activityinfo_activityreport "
                + query_condition )

    rows = cursor.fetchall()

    for row in rows:
        result[row[0]] = {
            'partner_id': row[0],
            'partner_label': row[1]
        }

    return result.values()


def get_governorates_list(ai_db, sections=None, partners=None, months=None, report_type=None):
    from django.db import connection

    result = {}
    cursor = connection.cursor()
    where_condition = ""
    funded_condition=""

    if ai_db.is_funded_by_unicef:
        funded_condition = " AND funded_by = 'UNICEF' "

    if sections:
        section_list = ", ".join("'" + str(n) + "'" for n in sections)
        where_condition += " AND reporting_section in (" + section_list + ")"
    if partners:
        partners_list = ", ".join("'" + str(n)+ "'" for n in partners)
        where_condition += " AND partner_id in (" + partners_list + ")"
    if months:
        month_list = ", ".join("'" + str(n) + "'" for n in months)
        where_condition += " AND date_part('month', start_date) in (" + month_list + ")"

    query_condition = "{}'{}'".format(" WHERE database_id =", str(ai_db.ai_id)) + funded_condition + where_condition

    if report_type == 'live':
        cursor.execute(
                "SELECT DISTINCT location_adminlevel_governorate_code, location_adminlevel_governorate "
                "FROM activityinfo_liveactivityreport "
                 + query_condition)
    else :
        cursor.execute(
                "SELECT DISTINCT location_adminlevel_governorate_code, location_adminlevel_governorate "
                "FROM activityinfo_activityreport "
                 + query_condition)

    rows = cursor.fetchall()

    for row in rows:
        result[row[0]] = {
            'location_adminlevel_governorate_code': row[0],
            'location_adminlevel_governorate': row[1]
        }

    return result.values()


def get_cadastrals_list(ai_db, report_type=None):
    from django.db import connection

    result = {}
    cursor = connection.cursor()

    if report_type == 'live':

        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT DISTINCT location_adminlevel_cadastral_area_code, location_adminlevel_cadastral_area "
                "FROM activityinfo_liveactivityreport "
                "WHERE database_id = %s AND funded_by = %s ",
                [str(ai_db.ai_id), 'UNICEF'])
        else:
            cursor.execute(
                "SELECT DISTINCT location_adminlevel_cadastral_area_code, location_adminlevel_cadastral_area "
                "FROM activityinfo_liveactivityreport "
                "WHERE database_id = %s ",
                [str(ai_db.ai_id)])
    else :
        if ai_db.is_funded_by_unicef:
            cursor.execute(
                "SELECT DISTINCT location_adminlevel_cadastral_area_code, location_adminlevel_cadastral_area "
                "FROM activityinfo_activityreport "
                "WHERE database_id = %s AND funded_by = %s ",
                [str(ai_db.ai_id), 'UNICEF'])
        else:
            cursor.execute(
                "SELECT DISTINCT location_adminlevel_cadastral_area_code, location_adminlevel_cadastral_area "
                "FROM activityinfo_activityreport "
                "WHERE database_id = %s ",
                [str(ai_db.ai_id)])

    rows = cursor.fetchall()

    for row in rows:
        result[row[0]] = {
            'location_adminlevel_cadastral_area_code': row[0],
            'location_adminlevel_cadastral_area': row[1]
        }

    return result.values()


def get_reporting_sections_list(ai_db,partners=None, govs=None,months=None, report_type=None):
    from django.db import connection

    result = {}
    cursor = connection.cursor()
    where_condition = ""
    funded_condition = ""

    if ai_db.is_funded_by_unicef:
        funded_condition = " AND funded_by = 'UNICEF' "

    if partners:
        partners_list = ", ".join("'" + str(n) + "'" for n in partners)
        where_condition += " AND partner_id in (" + partners_list + ")"
    if govs:
        govs_list = ", ".join("'" + str(n) + "'" for n in govs)
        where_condition = " and location_adminlevel_governorate_code in (" + govs_list + ")"
    if months:
        month_list = ", ".join("'" + str(n) + "'" for n in months)
        where_condition += " AND date_part('month', start_date) in (" + month_list + ")"

    query_condition = "{}'{}'".format(" WHERE database_id =", str(ai_db.ai_id)) + funded_condition + where_condition
    if report_type == 'live':
        cursor.execute(
                "SELECT DISTINCT reporting_section "
                "FROM activityinfo_liveactivityreport " + query_condition)
    else:
        cursor.execute(
                "SELECT DISTINCT reporting_section "
                "FROM activityinfo_activityreport "+query_condition
              )
    rows = cursor.fetchall()

    for row in rows:
        result[row[0]] = {
            'reporting_section': row[0],
        }

    return result.values()


def get_months_list(ai_db,partners=None, govs=None,sections=None, report_type=None):
    from django.db import connection

    result = {}
    months=[]
    cursor = connection.cursor()
    where_condition = ""
    funded_condition = ""

    if ai_db.is_funded_by_unicef:
        funded_condition = " AND funded_by = 'UNICEF' "

    if partners:
        partners_list = ", ".join("'" + str(n) + "'" for n in partners)
        where_condition += " AND partner_id in (" + partners_list + ")"
    if govs:
        govs_list = ", ".join("'" + str(n) + "'" for n in govs)
        where_condition = " and location_adminlevel_governorate_code in (" + govs_list + ")"
    if sections:
        section_list = ", ".join("'" + str(n) + "'" for n in sections)
        where_condition += " AND reporting_section in (" + section_list + ")"

    query_condition = "{}'{}'".format(" WHERE database_id =", str(ai_db.ai_id)) + funded_condition + where_condition
    if report_type == 'live':
        cursor.execute(
                "SELECT DISTINCT start_date "
                "FROM activityinfo_liveactivityreport " + query_condition

        )
    else:
        cursor.execute(
                "SELECT DISTINCT start_date "
                "FROM activityinfo_activityreport "+query_condition
              )
    rows = cursor.fetchall()

    for row in rows:
        if row[0] is not None:
            result[row[0]] = {
                'start_date': row[0],
            }
    sorted_list= sorted(result)
    for record in sorted_list:
        if record is not None:
            m = record.month
            if (m, calendar.month_name[m]) not in months:
                months.append((m, calendar.month_name[m]))

    return months


def calculate_indicators_highest_results(ai_db, report_type=None):
    from internos.activityinfo.models import Indicator, ActivityReport, LiveActivityReport
    from django.db import connection

    master_indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,is_cumulative=False)
    if master_indicators:
        linked_indicators = []
        for master_ind in master_indicators:
            indicators = Indicator.objects.filter(sub_indicators=master_ind.id).only(
                'id',
                'values',
                'values_gov',
                'values_partners',
                'values_partners_gov',
                'values_live',
                'values_gov_live',
                'values_partners_live',
                'values_partners_gov_live',
                'cumulative_values',
                'cumulative_values_live'
            )
            for ind in indicators:
                linked_indicators.append(ind)

        for ind in master_indicators:
            linked_indicators.append(ind)
        indicators_ids = []

        for ind in linked_indicators:
            value = ind.id
            indicators_ids.append(value)

        ids_condition = ','.join(("'{}'".format(str(n)) for n in indicators_ids))

        rows_data = {}
        cursor = connection.cursor()
        cursor.execute(
            "SELECT distinct ai.id, ai.ai_indicator, aa.id, aa.name, ai.values, ai.values_live, "
            "ai.values_gov, ai.values_gov_live, ai.values_partners, ai.values_partners_live, "
            "ai.values_partners_gov, ai.values_partners_gov_live "
            "FROM public.activityinfo_indicator ai, public.activityinfo_activity aa "
            "WHERE ai.activity_id = aa.id AND aa.database_id = %s and ai.id in(" + ids_condition +")",
            [ai_db.id])

        rows = cursor.fetchall()
        for row in rows:
            rows_data[row[0]] = row

        for indicator in linked_indicators:
            highest_month = {}
            highest_partner = {}
            highest_gov = {}
            highest_partner_gov = {}

            if indicator.id in rows_data:
                indicator_values = rows_data[indicator.id]

                if report_type == 'live':
                    values = indicator_values[5]  # values_live
                    values1 = indicator_values[7]  # values_gov_live
                    values2 = indicator_values[9]  # values_partners_live
                    values3 = indicator_values[11]  # values_partners_gov_live
                else:
                    values = indicator_values[4]  # values
                    values1 = indicator_values[6]  # values_gov
                    values2 = indicator_values[8]  # values_partners
                    values3 = indicator_values[10]  # values_partners_gov

                if len(values.values()) > 0:
                    highest_month = max(values.values())

                for key, value in values1.items():
                    keys = key.split('-')
                    gov = keys[1]
                    if gov in highest_gov:
                        if value > highest_gov[gov]:
                             highest_gov[gov] = value
                    else:
                        highest_gov[gov] = value

                for key, value in values2.items():
                    keys = key.split('-')
                    partner = keys[1]
                    if partner in highest_partner:
                       if value > highest_partner[partner]:
                           highest_partner[partner] = value
                    else:
                        highest_partner[partner] = value

                for key, value in values3.items():
                    keys = key.split('-')
                    gov_partner = '{}-{}'.format(keys[1], keys[2])
                    if gov_partner in highest_partner_gov:
                        if value > highest_partner_gov[gov_partner]:
                            highest_partner_gov[gov_partner] =  value
                    else:
                        highest_partner_gov[gov_partner] = value

                if report_type == 'live':
                    indicator.highest_values_live = {
                        'months': highest_month,
                        'partners': highest_partner,
                        'govs': highest_gov,
                        'partners_govs': highest_partner_gov
                    }
                else:
                    indicator.highest_values = {
                         'months': highest_month,
                         'partners': highest_partner,
                         'govs': highest_gov,
                         'partners_govs': highest_partner_gov
                    }
                indicator.save()

    # return linked_indicators.count()

