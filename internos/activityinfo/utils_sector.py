import os
import csv
import json
import datetime
import subprocess
from django.db.models import Sum, Q
from django.conf import settings


def reset_indicators_values(ai_id):
    from internos.activityinfo.models import Indicator

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_id)
    for indicator in indicators:
        indicator.values_sector = {}
        indicator.values_sites_sector = {}
        indicator.values_partners_sector = {}
        indicator.values_partners_sites_sector = {}
        indicator.cumulative_values_sector = {}
        indicator.save()

    return indicators.count()


def calculate_indicators_values(ai_db):
    print('reset_indicators_values_sector')
    reset_indicators_values(ai_db.ai_id)
    print('calculate_individual_indicators_values_sector')
    calculate_individual_indicators_values(ai_db)
    print('calculate_master_indicators_values')
    calculate_master_indicators_values(ai_db, True)
    print('calculate_master_indicators_values')
    calculate_master_indicators_values(ai_db)
    print('calculate_master_indicators_values_percentage')
    calculate_master_indicators_values_percentage(ai_db)
    print('calculate_master_indicators_values_denominator_multiplication')
    calculate_master_indicators_values_denominator_multiplication(ai_db)
    print('calculate_indicators_values_percentage')
    calculate_indicators_values_percentage(ai_db)
    print('calculate_indicators_cumulative_results')
    calculate_indicators_cumulative_results(ai_db)
    print('calculate_indicators_status')
    calculate_indicators_status(ai_db)

    return 0


def calculate_indicators_cumulative_results(ai_db):
    from django.db import connection
    from internos.activityinfo.models import Indicator

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id).only(
        'id',
        'values_sector',
        'values_sites_sector',
        'values_partners_sector',
        'values_partners_sites_sector',
        'cumulative_values_sector',
    )

    rows_data = {}
    cursor = connection.cursor()
    cursor.execute(
        "SELECT distinct ai.id, ai.ai_indicator, aa.id, aa.name, ai.values_sector, " 
        "ai.values_sites_sector, ai.values_partners_sector, " 
        "ai.values_partners_sites_sector "
        "FROM public.activityinfo_indicator ai, public.activityinfo_activity aa "
        "WHERE ai.activity_id = aa.id AND aa.database_id = %s",
        [ai_db.id])

    rows = cursor.fetchall()
    for row in rows:
        rows_data[row[0]] = row

    for indicator in indicators.iterator():
        values_month = {}
        values_partners = {}
        values_sites = {}
        values_partners_sites = {}
        if indicator.id in rows_data:
            indicator_values = rows_data[indicator.id]

            values = indicator_values[4]  # values_sector
            values1 = indicator_values[5]  # values_sites_sector
            values2 = indicator_values[6]  # values_partners
            values3 = indicator_values[7]  # values_partners_sites_sector

            c_value = 0
            for key, value in values.items():
                c_value += value
                values_month = c_value

            for key, value in values1.items():
                keys = key.split('-')
                site = keys[1]
                if site in values_sites:
                    values_sites[site] = values_sites[site] + value
                else:
                    values_sites[site] = value

            for key, value in values2.items():
                keys = key.split('-')
                partner = keys[1]
                if partner in values_partners:
                    values_partners[partner] = values_partners[partner] + value
                else:
                    values_partners[partner] = value

            for key, value in values3.items():
                keys = key.split('-')
                site_partner = '{}-{}'.format(keys[1], keys[2])
                if site_partner in values_partners_sites:
                    values_partners_sites[site_partner] = values_partners_sites[site_partner] + value
                else:
                    values_partners_sites[site_partner] = value

            indicator.cumulative_values_sector = {
                'months': values_month,
                'partners': values_partners,
                'sites': values_sites,
                'partners_sites': values_partners_sites
            }

            indicator.save()

    return indicators.count()


def calculate_indicators_tags():
    from internos.activityinfo.models import Indicator, IndicatorTag

    indicators = Indicator.objects.filter(Q(master_indicator=True) | Q(hpm_indicator=True))
    tags_gender = IndicatorTag.objects.filter(type='gender').only('id', 'name')
    tags_age = IndicatorTag.objects.filter(type='age').only('id', 'name')
    tags_nationality = IndicatorTag.objects.filter(type='nationality').only('id', 'name')
    tags_disability = IndicatorTag.objects.filter(type='disability').only('id', 'name')

    for indicator in indicators.iterator():
        m_value = 0
        try:
            m_value = indicator.cumulative_values_sector['months']
        except Exception:
            continue
        if isinstance(m_value, dict):
            m_value = 0
        sub_indicators = indicator.summation_sub_indicators.all().only(
            'values_tags_sector',
            'cumulative_values_sector',
        )
        for tag in tags_gender.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_gender_id=tag.id)

            value = 0
            for ind_tag in tag_sub_indicators:
                c_value = 0
                if 'months' in ind_tag.cumulative_values_sector:
                    c_value = ind_tag.cumulative_values_sector['months']

                if isinstance(c_value, dict):
                    c_value = 0

                value += float(c_value)

            try:
                indicator.values_tags_sector[tag.name] = float(value) * 100 / float(m_value)
            except Exception as ex:
                # print(ex.message)
                indicator.values_tags_sector[tag.name] = 0

        for tag in tags_age.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_age_id=tag.id)

            value = 0
            for ind_tag in tag_sub_indicators:
                c_value = 0
                if 'months' in ind_tag.cumulative_values_sector:
                    c_value = ind_tag.cumulative_values_sector['months']

                if isinstance(c_value, dict):
                    c_value = 0

                value += float(c_value)

            try:
                indicator.values_tags_sector[tag.name] = float(value) * 100 / float(m_value)
            except Exception as ex:
                # print(ex.message)
                indicator.values_tags_sector[tag.name] = 0

        for tag in tags_nationality.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_nationality_id=tag.id)

            value = 0
            for ind_tag in tag_sub_indicators:
                c_value = 0
                if 'months' in ind_tag.cumulative_values_sector:
                    c_value = ind_tag.cumulative_values_sector['months']

                if isinstance(c_value, dict):
                    c_value = 0

                value += float(c_value)

            try:
                indicator.values_tags_sector[tag.name] = float(value) * 100 / float(m_value)
            except Exception as ex:
                # print(ex.message)
                indicator.values_tags_sector[tag.name] = 0

        for tag in tags_disability.iterator():
            tag_sub_indicators = sub_indicators.filter(tag_disability_id=tag.id)

            value = 0
            for ind_tag in tag_sub_indicators:
                c_value = 0
                if 'months' in ind_tag.cumulative_values_sector:
                    c_value = ind_tag.cumulative_values_sector['months']

                if isinstance(c_value, dict):
                    c_value = 0

                value += float(c_value)

            try:
                indicator.values_tags_sector[tag.name] = float(value) * 100 / float(m_value)
            except Exception as ex:
                # print(ex.message)
                indicator.values_tags_sector[tag.name] = 0

        indicator.save()

    return indicators.count()


def calculate_master_indicators_values(ai_db, sub_indicators=False):
    from django.db import connection
    from internos.activityinfo.models import Indicator

    if sub_indicators:
        indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                              master_indicator_sub=True)
    else:
        indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                              master_indicator=True)

    indicators = indicators.only(
        'summation_sub_indicators',
        'values_sector',
        'values_sites_sector',
        'values_partners_sector',
        'values_partners_sites_sector',
    )

    rows_data = {}
    cursor = connection.cursor()
    cursor.execute(
        "SELECT distinct a1.id, a1.ai_indicator, ai.ai_indicator, ai.id, ai.values_sector, "
        "ai.values_sites_sector, ai.values_partners_sector, ai.values_partners_sites_sector, " 
        "a1.master_indicator, a1.master_indicator_sub "
        "FROM public.activityinfo_indicator a1, public.activityinfo_activity aa, " 
        "public.activityinfo_indicator_summation_sub_indicators ais, public.activityinfo_indicator ai " 
        "WHERE ai.activity_id = aa.id AND a1.id = ais.from_indicator_id AND ais.to_indicator_id = ai.id " 
        "AND aa.database_id = %s AND (a1.master_indicator = true or a1.master_indicator_sub = true)",
        [ai_db.id])

    rows = cursor.fetchall()
    for row in rows:
        if row[0] not in rows_data:
            rows_data[row[0]] = {}

        rows_data[row[0]][row[3]] = row

    for indicator in indicators.iterator():
        values_month = {}
        values_partners = {}
        values_sites_sector = {}
        values_partners_sites_sector = {}
        if indicator.id in rows_data:
            indicator_values = rows_data[indicator.id]
            for key, key_values in indicator_values.items():
                sub_indicator_values = indicator_values[key]

                values = sub_indicator_values[4]  # values
                values1 = sub_indicator_values[5]  # values_sites_sector
                values2 = sub_indicator_values[6]  # values_partners
                values3 = sub_indicator_values[7]  # values_partners_sites_sector

                for key in values:
                    val = values[key]
                    if key in values_month:
                        val = values_month[key] + val
                    values_month[key] = val

                for key in values1:
                    val = values1[key]
                    if key in values_sites_sector:
                        val = values_sites_sector[key] + val
                    values_sites_sector[key] = val

                for key in values2:
                    val = values2[key]
                    if key in values_partners:
                        val = values_partners[key] + val
                    values_partners[key] = val

                for key in values3:
                    val = values3[key]
                    if key in values_partners_sites_sector:
                        val = values_partners_sites_sector[key] + val
                    values_partners_sites_sector[key] = val

                indicator.values_sector = values_month
                indicator.values_sites_sector = values_sites_sector
                indicator.values_partners = values_partners
                indicator.values_partners_sites_sector = values_partners_sites_sector

            indicator.save()


def calculate_indicators_values_percentage(ai_db):
    from django.db import connection
    from internos.activityinfo.models import Indicator

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                          calculated_indicator=True).only(
        'denominator_indicator',
        'numerator_indicator',
        'denominator_multiplication',
        'values_sector',
        'values_sites_sector',
        'values_partners_sector',
        'values_partners_sites_sector',
    )

    rows_data = {}
    cursor = connection.cursor()
    cursor.execute(
        "SELECT distinct a1.id, a1.calculated_percentage, aa.name, ai.id, ai.values_sector, "
        "ai.values_sites_sector, ai.values_partners_sector, ai.values_partners_sites_sector, "
        "FROM public.activityinfo_indicator a1, public.activityinfo_activity aa, "
        "public.activityinfo_indicator_sub_indicators ais, public.activityinfo_indicator ai "
        "WHERE ai.activity_id = aa.id AND a1.id = ais.from_indicator_id AND ais.to_indicator_id = ai.id " 
        "AND aa.database_id = %s AND a1.calculated_indicator = true",
        [ai_db.id])

    rows = cursor.fetchall()
    for row in rows:
        rows_data[row[0]] = row

    for indicator in indicators.iterator():
        values_month = {}
        values_partners = {}
        values_sites_sector = {}
        values_partners_sites_sector = {}
        if indicator.id in rows_data:
            indicator_values = rows_data[indicator.id]
            reporting_level = indicator_values[2]
            calculated_percentage = indicator_values[1]

            values = indicator_values[4]  # values
            values1 = indicator_values[5]  # values_sites_sector
            values2 = indicator_values[6]  # values_partners
            values3 = indicator_values[7]  # values_partners_sites_sector

            for key in values:
                val = values[key]
                try:
                    if reporting_level == 'Municipality level':
                        values_month[key] = val * calculated_percentage / 100
                    elif reporting_level == 'Site level':
                        values_month[key] = val * calculated_percentage / 100
                except Exception as ex:
                    values_month[key] = 0

            for key in values1:
                val = values1[key]
                try:
                    if reporting_level == 'Municipality level':
                        values_sites_sector[key] = val * calculated_percentage / 100
                    elif reporting_level == 'Site level':
                        values_sites_sector[key] = val * calculated_percentage / 100
                except Exception as ex:
                    values_sites_sector[key] = 0

            for key in values2:
                val = values2[key]
                try:
                    if reporting_level == 'Municipality level':
                        values_partners[key] = val * calculated_percentage / 100
                    elif reporting_level == 'Site level':
                        values_partners[key] = val * calculated_percentage / 100
                except Exception as ex:
                    values_partners[key] = 0

            for key in values3:
                val = values3[key]
                try:
                    if reporting_level == 'Municipality level':
                        values_partners_sites_sector[key] = val * calculated_percentage / 100
                    elif reporting_level == 'Site level':
                        values_partners_sites_sector[key] = val * calculated_percentage / 100
                except Exception as ex:
                    values_partners_sites_sector[key] = 0

            indicator.values_sector = values_month
            indicator.values_sites_sector = values_sites_sector
            indicator.values_partners = values_partners
            indicator.values_partners_sites_sector = values_partners_sites_sector

            indicator.save()


def calculate_master_indicators_values_percentage(ai_db):
    from internos.activityinfo.models import Indicator, ActivityReport

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                          master_indicator=True,
                                          measurement_type='percentage').only(
        'denominator_indicator',
        'numerator_indicator',
        'denominator_multiplication',
        'values_sector',
        'values_sites_sector',
        'values_partners_sector',
        'values_partners_sites_sector',
    )
    last_month = int(datetime.datetime.now().strftime("%m"))

    report = ActivityReport.objects.filter(database_id=ai_db.ai_id)
    report = report.only('partner_id', 'location_adminlevel_cadastral_area_code')

    partners = report.values('partner_id').distinct()
    sites = report.values('location_adminlevel_cadastral_area_code').distinct()
    sites1 = report.values('location_adminlevel_cadastral_area_code').distinct()

    for indicator in indicators.iterator():
        for month in range(1, last_month):
            month = str(month)
            values_sites_sector = {}
            values_partners = {}
            values_partners_sites_sector = {}
            denominator_indicator = indicator.denominator_indicator
            numerator_indicator = indicator.numerator_indicator
            if not denominator_indicator or not numerator_indicator:
                continue
            try:
                denominator = denominator_indicator.values_sector[month] if month in denominator_indicator.values_sector else 0
                numerator = numerator_indicator.values_sector[month] if month in numerator_indicator.values_sector else 0
                values_month = numerator / denominator
            except Exception:
                values_month = 0

            for site1 in sites1:
                key = "{}-{}".format(month, site1['location_adminlevel_cadastral_area_code'])
                try:
                    denominator = denominator_indicator.values_sites_sector[key] if key in denominator_indicator.values_sites_sector else 0
                    numerator = numerator_indicator.values_sites_sector[key] if key in numerator_indicator.values_sites_sector else 0
                    values_sites_sector[key] = numerator / denominator
                except Exception:
                    values_sites_sector[key] = 0

            for partner in partners:
                key1 = "{}-{}".format(month, partner['partner_id'])

                try:
                    denominator = denominator_indicator.values_partners_sector[key1] if key1 in denominator_indicator.values_partners_sector else 0
                    numerator = numerator_indicator.values_partners_sector[key1] if key1 in numerator_indicator.values_partners_sector else 0
                    values_partners[key1] = numerator / denominator
                except Exception:
                    values_partners[key1] = 0

                for site in sites:
                    key2 = "{}-{}-{}".format(month, partner['partner_id'], site['location_adminlevel_cadastral_area_code'])
                    try:
                        denominator = denominator_indicator.values_partners_sites_sector[key2] if key2 in denominator_indicator.values_partners_sites_sector else 0
                        numerator = numerator_indicator.values_partners_sites_sector[key2] if key2 in numerator_indicator.values_partners_sites_sector else 0
                        values_partners_sites_sector[key2] = numerator / denominator
                    except Exception:
                        values_partners_sites_sector[key2] = 0

            indicator.values_sector[month] = values_month
            indicator.values_sites_sector.update(values_sites_sector)
            indicator.values_partners_sector.update(values_partners)
            indicator.values_partners_sites_sector.update(values_partners_sites_sector)

        indicator.save()


def calculate_master_indicators_values_denominator_multiplication(ai_db):
    from internos.activityinfo.models import Indicator, ActivityReport

    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id,
                                          master_indicator=True,
                                          measurement_type='percentage_x').only(
        'denominator_indicator',
        'numerator_indicator',
        'denominator_multiplication',
        'values_sector',
        'values_sites_sector',
        'values_partners_sector',
        'values_partners_sites_sector',
    )

    last_month = int(datetime.datetime.now().strftime("%m"))

    report = ActivityReport.objects.filter(database_id=ai_db.ai_id)
    report = report.only('partner_id', 'location_adminlevel_cadastral_area_code')

    partners = report.values('partner_id').distinct()
    sites = report.values('location_adminlevel_cadastral_area_code').distinct()
    sites1 = report.values('location_adminlevel_cadastral_area_code').distinct()

    for indicator in indicators.iterator():
        for month in range(1, last_month):
            month = str(month)
            values_sites_sector = {}
            values_partners = {}
            values_partners_sites_sector = {}
            denominator_indicator = indicator.denominator_indicator
            numerator_indicator = indicator.numerator_indicator
            denominator_multiplication = indicator.denominator_multiplication
            if not denominator_indicator or not numerator_indicator:
                continue
            try:
                denominator = denominator_indicator.values_sector[month] if month in denominator_indicator.values_sector else 0
                numerator = numerator_indicator.values_sector[month] if month in numerator_indicator.values_sector else 0
                denominator = denominator * denominator_multiplication
                values_month = numerator / denominator
            except Exception as ex:
                values_month = 0

            for site1 in sites1:
                key = "{}-{}".format(month, site1['location_adminlevel_cadastral_area_code'])
                try:
                    denominator = denominator_indicator.values_sites_sector[key] if key in denominator_indicator.values_sites_sector else 0
                    numerator = numerator_indicator.values_sites_sector[key] if key in numerator_indicator.values_sites_sector else 0
                    denominator = denominator * denominator_multiplication
                    values_sites_sector[key] = numerator / denominator
                except Exception:
                    values_sites_sector[key] = 0

            for partner in partners:
                key1 = "{}-{}".format(month, partner['partner_id'])

                try:
                    denominator = denominator_indicator.values_partners_sector[key1] if key1 in denominator_indicator.values_partners_sector else 0
                    numerator = numerator_indicator.values_partners_sector[key1] if key1 in numerator_indicator.values_partners_sector else 0
                    denominator = denominator * denominator_multiplication
                    values_partners[key1] = numerator / denominator
                except Exception:
                    values_partners[key1] = 0

                for site in sites:
                    key2 = "{}-{}-{}".format(month, partner['partner_id'], site['location_adminlevel_cadastral_area_code'])
                    try:
                        denominator = denominator_indicator.values_partners_sites_sector[key2] if key2 in denominator_indicator.values_partners_sites_sector else 0
                        numerator = numerator_indicator.values_partners_sites_sector[key2] if key2 in numerator_indicator.values_partners_sites_sector else 0
                        denominator = denominator * denominator_multiplication
                        values_partners_sites_sector[key2] = numerator / denominator
                    except Exception:
                        values_partners_sites_sector[key2] = 0

            indicator.values_sector[month] = values_month
            indicator.values_sites_sector.update(values_sites_sector)
            indicator.values_partners_sector.update(values_partners)
            indicator.values_partners_sites_sector.update(values_partners_sites_sector)

        indicator.save()


def calculate_individual_indicators_values(ai_db):
    from django.db import connection
    from internos.activityinfo.models import Indicator

    last_month = int(datetime.datetime.now().strftime("%m"))
    ai_id = str(ai_db.ai_id)
    indicators = Indicator.objects.filter(activity__database__ai_id=ai_db.ai_id).exclude(ai_id__isnull=True).only(
        'ai_indicator',
        'values_sector',
        'values_sites_sector',
        'values_partners_sector',
        'values_partners_sites_sector')

    rows_months = {}
    rows_partners = {}
    rows_sites = {}
    rows_partners_sites = {}
    cursor = connection.cursor()
    for month in range(1, last_month):
        month = str(month)
        cursor.execute(
            "SELECT indicator_id, SUM(indicator_value) as indicator_value "
            "FROM activityinfo_activityreport "
            "WHERE date_part('month', start_date) = %s AND database_id = %s "
            "GROUP BY indicator_id",
            [month, ai_id])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_months:
                rows_months[row[0]] = {}
            rows_months[row[0]][month] = row[1]

        cursor.execute(
            "SELECT indicator_id, SUM(indicator_value) as indicator_value, location_adminlevel_cadastral_area_code "
            "FROM activityinfo_activityreport "
            "WHERE date_part('month', start_date) = %s AND database_id = %s "
            "GROUP BY indicator_id, location_adminlevel_cadastral_area_code",
            [month, ai_id])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_sites:
                rows_sites[row[0]] = {}
            key = "{}-{}".format(month, row[2])
            rows_sites[row[0]][key] = row[1]

        cursor.execute(
            "SELECT indicator_id, SUM(indicator_value) as indicator_value, partner_id "
            "FROM activityinfo_activityreport "
            "WHERE date_part('month', start_date) = %s AND database_id = %s "
            "GROUP BY indicator_id, partner_id",
            [month, ai_id])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_partners:
                rows_partners[row[0]] = {}
            key = "{}-{}".format(month, row[2])
            rows_partners[row[0]][key] = row[1]

        cursor.execute(
            "SELECT indicator_id, SUM(indicator_value) as indicator_value, location_adminlevel_cadastral_area_code, partner_id "
            "FROM activityinfo_activityreport "
            "WHERE date_part('month', start_date) = %s AND database_id = %s "
            "GROUP BY indicator_id, location_adminlevel_cadastral_area_code, partner_id",
            [month, ai_id])
        rows = cursor.fetchall()

        for row in rows:
            if row[0] not in rows_partners_sites:
                rows_partners_sites[row[0]] = {}
            key = "{}-{}-{}".format(month, row[2], row[3])
            rows_partners_sites[row[0]][key] = row[1]

    for indicator in indicators.iterator():
        if indicator.ai_indicator in rows_months:
            indicator.values_sector = rows_months[indicator.ai_indicator]

        if indicator.ai_indicator in rows_partners:
            indicator.values_partners_sector = rows_partners[indicator.ai_indicator]

        if indicator.ai_indicator in rows_sites:
            indicator.values_sites_sector = rows_sites[indicator.ai_indicator]

        if indicator.ai_indicator in rows_partners_sites:
            indicator.values_partners_sites_sector = rows_partners_sites[indicator.ai_indicator]

        indicator.save()


def calculate_indicators_status(database):
    from internos.activityinfo.models import Indicator

    year_days = 365
    today = datetime.datetime.now()
    reporting_year = database.reporting_year
    beginning_year = datetime.datetime(int(reporting_year.name), 01, 01)
    delta = today - beginning_year
    total_days = delta.days + 1
    days_passed_per = (total_days * 100) / year_days

    indicators = Indicator.objects.filter(activity__database__ai_id=database.ai_id)

    for indicator in indicators:
        cumulative_per = indicator.cumulative_per_sector
        off_track = days_passed_per - 10
        over_target = days_passed_per + 10
        if cumulative_per < off_track:
            indicator.status_sector = 'Off Track'
            indicator.status_color_sector = '#FF0000'
        elif cumulative_per > over_target:
            indicator.status_sector = 'Over Target'
            indicator.status_color_sector = '#FFA500'
        else:
            indicator.status_sector = 'On Track'
            indicator.status_color_sector = '#008000'

        indicator.save()

    return indicators.count()


def update_hpm_table_docx(indicators, month, month_name, filename):

    from docx import Document
    from internos.activityinfo.templatetags.util_tags import get_indicator_hpm_data

    path = os.path.dirname(os.path.abspath(__file__))
    path2file = path+'/AIReports/HPM Table template 2019.docx'

    document = Document(path2file)

    document.paragraphs[0].runs[1].text = month_name

    # Education 1
    document.tables[0].rows[2].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3942, month)['cumulative'])
    document.tables[0].rows[2].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3942, month)['last_report_changes'])

    document.tables[0].rows[3].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3942, month)['cumulative'])
    document.tables[0].rows[3].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3942, month)['last_report_changes'])

    document.tables[0].rows[4].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3942, month)['cumulative'])
    document.tables[0].rows[4].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3942, month)['last_report_changes'])

    # Education 2
    document.tables[0].rows[5].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3985, month)['cumulative'])
    document.tables[0].rows[5].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3985, month)['last_report_changes'])

    document.tables[0].rows[6].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3985, month)['cumulative'])
    document.tables[0].rows[6].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3985, month)['last_report_changes'])

    document.tables[0].rows[7].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3985, month)['cumulative'])
    document.tables[0].rows[7].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3985, month)['last_report_changes'])

    document.tables[0].rows[8].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3985, month)['cumulative'])
    document.tables[0].rows[8].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3985, month)['last_report_changes'])

    # CP
    document.tables[0].rows[10].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1789, month)['cumulative'])
    document.tables[0].rows[10].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1789, month)['last_report_changes'])

    document.tables[0].rows[11].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1654, month)['cumulative'])
    document.tables[0].rows[11].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1654, month)['last_report_changes'])

    document.tables[0].rows[12].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(2740, month)['cumulative'])
    document.tables[0].rows[12].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(2740, month)['last_report_changes'])

    # WASH
    document.tables[0].rows[14].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1519, month)['cumulative'])
    document.tables[0].rows[14].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1519, month)['last_report_changes'])

    document.tables[0].rows[15].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1527, month)['cumulative'])
    document.tables[0].rows[15].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1527, month)['last_report_changes'])

    document.tables[0].rows[16].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1504, month)['cumulative'])
    document.tables[0].rows[16].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1504, month)['last_report_changes'])

    document.tables[0].rows[17].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1535, month)['cumulative'])
    document.tables[0].rows[17].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1535, month)['last_report_changes'])

    # H&N
    document.tables[0].rows[19].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(2621, month)['cumulative'])
    document.tables[0].rows[19].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(2621, month)['last_report_changes'])

    document.tables[0].rows[20].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(2638, month)['cumulative'])
    document.tables[0].rows[20].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(2638, month)['last_report_changes'])

    document.tables[0].rows[21].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(2178, month)['cumulative'])
    document.tables[0].rows[21].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(2178, month)['last_report_changes'])

    # Y&A
    document.tables[0].rows[23].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3032, month)['cumulative'])
    document.tables[0].rows[23].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3032, month)['last_report_changes'])

    document.tables[0].rows[24].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3071, month)['cumulative'])
    document.tables[0].rows[24].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3071, month)['last_report_changes'])

    document.tables[0].rows[25].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3179, month)['cumulative'])
    document.tables[0].rows[25].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3179, month)['last_report_changes'])

    document.tables[0].rows[26].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3347, month)['cumulative'])
    document.tables[0].rows[26].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(3347, month)['last_report_changes'])

    # SP
    document.tables[0].rows[28].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(850, month)['cumulative'])
    document.tables[0].rows[28].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(850, month)['last_report_changes'])

    document.tables[0].rows[29].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(783, month)['cumulative'])
    document.tables[0].rows[29].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(783, month)['last_report_changes'])

    # C4D
    document.tables[0].rows[31].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1396, month)['cumulative'])
    document.tables[0].rows[31].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1396, month)['last_report_changes'])

    # PPL
    document.tables[0].rows[33].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1121, month)['cumulative'])
    document.tables[0].rows[33].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1121, month)['last_report_changes'])

    document.tables[0].rows[34].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1149, month)['cumulative'])
    document.tables[0].rows[34].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1149, month)['last_report_changes'])

    document.tables[0].rows[35].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1178, month)['cumulative'])
    document.tables[0].rows[35].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1178, month)['last_report_changes'])

    document.tables[0].rows[36].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1195, month)['cumulative'])
    document.tables[0].rows[36].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1195, month)['last_report_changes'])

    document.tables[0].rows[37].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1293, month)['cumulative'])
    document.tables[0].rows[37].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1293, month)['last_report_changes'])

    document.tables[0].rows[38].cells[6].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1299, month)['cumulative'])
    document.tables[0].rows[38].cells[7].paragraphs[0].runs[0].text = str(get_indicator_hpm_data(1299, month)['last_report_changes'])

    # Footnotes

    #  CP
    hpm_data = get_indicator_hpm_data(1789, month)
    document.tables[0].rows[42].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Boys', hpm_data['male'])
    document.tables[0].rows[42].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Girls', hpm_data['female'])

    hpm_data = get_indicator_hpm_data(1654, month)
    document.tables[0].rows[43].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Boys', hpm_data['male'])
    document.tables[0].rows[43].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Girls', hpm_data['female'])

    #  WASH
    hpm_data = get_indicator_hpm_data(1519, month)
    document.tables[0].rows[44].cells[4].paragraphs[0].runs[0].text = '{}: {}%'.format('Boys', hpm_data['boys'])
    document.tables[0].rows[44].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Girls', hpm_data['girls'])
    document.tables[0].rows[44].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Female', hpm_data['female'])
    document.tables[0].rows[44].cells[9].paragraphs[0].runs[0].text = '{}: {}%'.format('Male', hpm_data['male'])

    hpm_data = get_indicator_hpm_data(1527, month)
    document.tables[0].rows[45].cells[4].paragraphs[0].runs[0].text = '{}: {}%'.format('Boys', hpm_data['boys'])
    document.tables[0].rows[45].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Girls', hpm_data['girls'])
    document.tables[0].rows[45].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Female', hpm_data['female'])
    document.tables[0].rows[45].cells[9].paragraphs[0].runs[0].text = '{}: {}%'.format('Male', hpm_data['male'])

    # hpm_data = get_indicator_hpm_data(1504, month)
    # document.tables[0].rows[46].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Boys', hpm_data['male'])
    # document.tables[0].rows[46].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Girls', hpm_data['female'])

    hpm_data = get_indicator_hpm_data(1535, month)
    document.tables[0].rows[47].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Boys', hpm_data['male'])
    document.tables[0].rows[47].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Girls', hpm_data['female'])

    #  Y&A
    hpm_data = get_indicator_hpm_data(3032, month)
    document.tables[0].rows[49].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Male', hpm_data['male'])
    document.tables[0].rows[49].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Female', hpm_data['female'])

    hpm_data = get_indicator_hpm_data(3071, month)
    document.tables[0].rows[50].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Male', hpm_data['male'])
    document.tables[0].rows[50].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Female', hpm_data['female'])

    hpm_data = get_indicator_hpm_data(3179, month)
    document.tables[0].rows[51].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Male', hpm_data['male'])
    document.tables[0].rows[51].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Female', hpm_data['female'])

    hpm_data = get_indicator_hpm_data(3347, month)
    document.tables[0].rows[52].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Male', hpm_data['male'])
    document.tables[0].rows[52].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Female', hpm_data['female'])

    #  PPL
    hpm_data = get_indicator_hpm_data(1121, month)
    document.tables[0].rows[55].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Boys', hpm_data['male'])
    document.tables[0].rows[55].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Girls', hpm_data['female'])

    hpm_data = get_indicator_hpm_data(1149, month)
    document.tables[0].rows[56].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Boys', hpm_data['male'])
    document.tables[0].rows[56].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Girls', hpm_data['female'])

    hpm_data = get_indicator_hpm_data(1178, month)
    document.tables[0].rows[57].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Boys', hpm_data['male'])
    document.tables[0].rows[57].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Girls', hpm_data['female'])

    hpm_data = get_indicator_hpm_data(1195, month)
    document.tables[0].rows[58].cells[6].paragraphs[0].runs[0].text = '{}: {}%'.format('Boys', hpm_data['male'])
    document.tables[0].rows[58].cells[7].paragraphs[0].runs[0].text = '{}: {}%'.format('Girls', hpm_data['female'])

    path2file2 = '{}/{}/{}'.format(path, 'AIReports', filename)
    document.save(path2file2)
    return path2file2

