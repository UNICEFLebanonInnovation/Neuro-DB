__author__ = 'achamseddine'

import os
import csv
import time
import json


def read_docx():

    # import sys
    # import docx
    from docx import Document

    path = os.path.dirname(os.path.abspath(__file__))
    path2file = path+'/HACT_TDH-it_Arsal_Center_100419.docx'

    document = Document(path2file)
    for paragraph in document.paragraphs:
        p = paragraph._element
        # print(p.xml)
        checkBoxes = p.xpath('//w:sym')
        # print(checkBoxes)
        # if checkBoxes:
        #     print(p.xml)
        #     break
        # for item in checkBoxes:
        #     print(item.values())
    document = Document(path2file)

    table = document.tables[1]
    # print(document.tables[0].rows[0].cells[0].text)
    # print(document.tables[1].rows[0].cells[0].text)
    # print(document.tables[1].rows[2].cells[0].text)
    # print(document.tables[1].rows[3].cells[0].text)
    # print(document.tables[1].rows[4].cells[0].text)
    # print(document.tables[1].rows[5].cells[0].text)
    # print(document.tables[1].rows[6].cells[0].text)
    # print(document.tables[1].rows[6].cells[0].paragraphs[1].runs[0].text)
    # print(document.tables[1].rows[6].cells[0].paragraphs[1].runs[1].text)
    # print(document.tables[1].rows[6].cells[0].paragraphs[1].runs[2].text)
    # print(document.tables[1].rows[6].cells[0].paragraphs[1].runs[3].text)
    #
    # print(document.tables[1].rows[6].cells[0].xml)
    # print(document.tables[1].rows[6].cells[0].paragraphs[8]._element.xml)
    # print(document.tables[1].rows[6].cells[0].paragraphs[8].runs[0].text)
    # print(document.tables[1].rows[6].cells[0].paragraphs[8].runs[1].text)
    # print(document.tables[1].rows[6].cells[0].paragraphs[8].runs[2].text)
    # print(document.tables[1].rows[6].cells[0].paragraphs[8].text)
    # print(document.tables[1].rows[6].cells[0]._element.xml)

    paragraph = document.tables[1].rows[6].cells[0].paragraphs[7]
    print(paragraph.text)
    element = paragraph._element
    items = element.xpath('.//w:sym')
    for item in items:
        print(item.values())

    paragraph = document.tables[1].rows[6].cells[0].paragraphs[8]
    print(paragraph.text)
    element = paragraph._element
    items = element.xpath('.//w:sym')
    for item in items:
        print(item.values())

    # print(document.tables[1].rows[16].cells[0].paragraphs[0].text)
    # print(document.tables[1].rows[17].cells[0].paragraphs[0].text)
    # print(document.tables[1].rows[18].cells[0].paragraphs[0].text)
    print(document.tables[1].rows[19].cells[0].paragraphs[0].text)

    # print(document.tables[1].rows[6].cells[0].paragraphs[1].runs[4].text)
    # print(document.tables[1].rows[6].cells[0].paragraphs[1].runs[5].text)
    # print(document.tables[1].rows[6].cells[0].paragraphs[1].runs[6].text)
    # print(document.tables[1].rows[6].cells[0].paragraphs[1].runs[7].text)
    # print(document.tables[1].rows[6].cells[0].paragraphs[1]._element.xml)

